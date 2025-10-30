#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# use of main8.py

"""
Enhanced Site Survey Generator with Image Support
เพิ่มฟีเจอร์การแทรกรูปภาพใน More Detail section
"""

import os
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ImageSupportedSiteSurveyGenerator:
    """Site Survey Generator พร้อม Image Support"""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path or 'site survey.docx'
        self.image_folder = 'site_survey_images'
        
        # สร้าง folder สำหรับเก็บรูป
        os.makedirs(self.image_folder, exist_ok=True)
    
    def add_image_to_cell(self, cell, image_path: str, max_width_inches: float = 2.0, max_height_inches: float = 3.0):
        """
        แทรกรูปภาพลงใน cell ของตาราง โดยรักษาสัดส่วนและไม่เกินขนาดที่กำหนด
        
        Args:
            cell: Cell object จาก python-docx
            image_path: path ของไฟล์รูปภาพ
            max_width_inches: ความกว้างสูงสุดของรูป (inches)
            max_height_inches: ความสูงสูงสุดของรูป (inches)
        """
        try:
            if not os.path.exists(image_path):
                print(f"Warning: Image not found: {image_path}")
                return False
            
            # ลบข้อความเดิมใน cell
            cell.text = ""
            
            # เพิ่ม paragraph ใหม่
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            
            # กำหนด alignment เป็นตรงกลาง
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # อ่านขนาดรูปภาพจริง
            from PIL import Image
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                
            # คำนวณ aspect ratio
            aspect_ratio = img_width / img_height
            
            # คำนวณขนาดที่เหมาะสมโดยรักษาสัดส่วน
            if img_width / max_width_inches > img_height / max_height_inches:
                # จำกัดด้วยความกว้าง
                final_width = max_width_inches
                final_height = final_width / aspect_ratio
            else:
                # จำกัดด้วยความสูง
                final_height = max_height_inches
                final_width = final_height * aspect_ratio
            
            # เพิ่มรูปภาพด้วยขนาดที่คำนวณได้
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(final_width), height=Inches(final_height))
            
            print(f"✅ Successfully added image to cell: {image_path} (Size: {final_width:.2f}\" x {final_height:.2f}\")")
            return True
            
        except Exception as e:
            print(f"❌ Error adding image: {str(e)}")
            return False
    
    def add_multiple_images_to_cell(self, cell, image_paths: List[str], 
                                width_inches: float = 1.3, 
                                max_images_per_row: int = 2):
        """
        แทรกหลายรูปภาพลงใน cell (จัดเรียงแบบ grid)
        """
        try:
            if not image_paths:
                return False
            
            # ลบข้อความเดิม
            cell.text = ""
            
            # เพิ่มรูปภาพทีละรูป
            images_added = 0
            current_paragraph = None
            
            for i, image_path in enumerate(image_paths):
                if not os.path.exists(image_path):
                    print(f"Warning: Image not found: {image_path}")
                    continue
                
                # สร้าง paragraph ใหม่สำหรับแต่ละแถว
                if i % max_images_per_row == 0:
                    if current_paragraph:
                        # เพิ่มระยะห่างระหว่างแถว
                        current_paragraph.add_run().add_break()
                    current_paragraph = cell.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # เพิ่มรูปภาพ
                run = current_paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                
                # เพิ่มช่องว่างระหว่างรูป (ยกเว้นรูปสุดท้ายของแถว)
                if (i + 1) % max_images_per_row != 0 and i < len(image_paths) - 1:
                    run.add_text("  ")
                
                images_added += 1
            
            print(f"✅ Added {images_added} images to cell")
            return images_added > 0
            
        except Exception as e:
            print(f"❌ Error adding multiple images: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def fill_template_with_images(self, doc, product: Dict, 
                                image_mapping: Dict[str, List[str]]) -> bool:
        """
        เติมข้อมูลและรูปภาพลงใน template
        หารูปจากเซลล์ที่มีข้อความ "image"
        """
        try:
            if not doc.tables:
                print("No tables found in template")
                return False
            
            table = doc.tables[0]
            ref = product.get('ref', '')
            
            print(f"\n🔍 Processing ref: {ref}")
            print(f"   Available images for this ref: {len(image_mapping.get(ref, []))}")
            
            # หาเซลล์ที่มีข้อความ "image"
            image_cell_found = False
            
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().lower()
                    
                    # ถ้าเจอเซลล์ที่มีข้อความ "image" 
                    if cell_text == 'image':
                        print(f"   Found 'image' cell at row {row_index}, cell {cell_index}")
                        
                        # ดึงรูปภาพสำหรับ ref นี้
                        images = image_mapping.get(ref, [])
                        
                        if images:
                            print(f"   Replacing 'image' text with {len(images)} actual images")
                            
                            # ลบข้อความ "image"
                            cell.text = ""
                            
                            # เพิ่มรูปภาพ
                            success = self.add_multiple_images_to_cell(
                                cell, 
                                images, 
                                width_inches=1.3,  # ขนาดเล็กลงเพื่อให้พอดีกับเซลล์
                                max_images_per_row=2
                            )
                            
                            if success:
                                image_cell_found = True
                                print(f"✅ Successfully replaced 'image' with {len(images)} photos for ref: {ref}")
                            else:
                                print(f"❌ Failed to add images for ref: {ref}")
                        else:
                            print(f"ℹ️ No images found for ref: {ref}, keeping 'image' text")
                        
                        break
                
                if image_cell_found:
                    break
            
            if not image_cell_found:
                print(f"⚠️ No 'image' cell found in template")
            
            return image_cell_found
            
        except Exception as e:
            print(f"❌ Error filling template with images: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def organize_images_by_ref(self, ocr_results: List[Dict]) -> Dict[str, List[str]]:
        """
        จัดระเบียบรูปภาพตาม ref code จากข้อมูล OCR
        ✅ IMPROVED: Exact matching only, no fuzzy matching
        """
        image_mapping = {}
        
        for ocr_result in ocr_results:
            if not ocr_result.get('success'):
                continue
            
            # ดึง ref และ image path
            text = ocr_result.get('text', '')
            image_path = ocr_result.get('image_path') or ocr_result.get('filename')
            
            # แยก ref จากข้อความ
            import re
            ref = None
            
            # ✅ IMPROVED: More precise regex that captures full ref including decimals
            ref_match = re.search(r'\bref[:\s]*([DW][A-Z]?\d+(?:\.\d+)?)\b', text, re.IGNORECASE)
            if ref_match:
                ref = ref_match.group(1).upper()
            
            # ถ้าไม่เจอ ลองดูจาก ocr_result โดยตรง
            if not ref and 'ref' in ocr_result:
                ref = ocr_result.get('ref', '').upper()
            
            # ✅ IMPROVED: Try exact filename matching with full ref pattern
            if not ref and image_path:
                filename = os.path.basename(image_path)
                # Try to extract ref from filename with decimal support
                filename_match = re.match(r'^([DW][A-Z]?\d+(?:\.\d+)?)\b', filename, re.IGNORECASE)
                if filename_match:
                    ref = filename_match.group(1).upper()
            
            if ref and image_path and os.path.exists(image_path):
                if ref not in image_mapping:
                    image_mapping[ref] = []
                
                image_mapping[ref].append(image_path)
                print(f"✅ Mapped image for ref {ref}: {image_path}")
            else:
                print(f"⚠️ Could not map image: ref={ref}, path={image_path}, exists={os.path.exists(image_path) if image_path else False}")
        
        return image_mapping
    
    def generate_site_survey_with_images(self, products: List[Dict], 
                                        ocr_results: List[Dict],
                                        output_path: str) -> Dict[str, Any]:
        """
        สร้าง Site Survey พร้อมรูปภาพ
        
        Args:
            products: list ของ product data
            ocr_results: list ของ OCR results (มี image path)
            output_path: path สำหรับบันทึกไฟล์
        """
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not available")
            
            # จัดระเบียบรูปภาพ
            image_mapping = self.organize_images_by_ref(ocr_results)
            print(f"\n📸 Image mapping created for {len(image_mapping)} refs")
            
            # สร้างไฟล์ทีละ product
            import tempfile
            temp_files = []
            
            for i, product in enumerate(products):
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                temp_files.append(temp_file.name)
                temp_file.close()
                
                try:
                    # โหลด template
                    doc = Document(self.template_path)
                    
                    # เติมข้อมูลปกติ
                    self._fill_basic_data(doc, product)
                    
                    # เติมรูปภาพ
                    self.fill_template_with_images(doc, product, image_mapping)
                    
                    # บันทึก
                    doc.save(temp_file.name)
                    print(f"✅ Created page {i+1} for {product.get('ref')}")
                    
                except Exception as e:
                    print(f"❌ Error creating page {i+1}: {e}")
                    continue
            
            # รวมไฟล์
            if temp_files:
                self._merge_files(temp_files, output_path)
                
                # ลบไฟล์ชั่วคราว
                for temp_file in temp_files:
                    try:
                        os.unlink(temp_file)
                    except:
                        pass
                
                return {
                    'success': True,
                    'file_path': output_path,
                    'message': f'สร้าง Site Survey พร้อมรูปภาพสำเร็จ ({len(products)} หน้า)',
                    'images_added': sum(len(imgs) for imgs in image_mapping.values())
                }
            else:
                raise Exception("ไม่สามารถสร้างไฟล์ได้")
            
        except Exception as e:
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }
    
    def _fill_basic_data(self, doc, product: Dict):
        """เติมข้อมูลพื้นฐาน (ไม่รวมรูปภาพ)"""
        if not doc.tables:
            return
        
        table = doc.tables[0]
        current_date = datetime.now().strftime('%Y-%m-%d')
        
        replacements = {
            'ref1': product.get('ref', ''),
            'Series1': product.get('series', ''),
            'product_type_main': product.get('product_type', ''),
            'color1': product.get('color', ''),  # เปลี่ยนจาก Color1 เป็น color1
            'Glass1': product.get('glass', ''),
            'Screen1': product.get('insect_screen', 'No'),
            'W1': str(product.get('width', 0)),
            'H1': str(product.get('height', 0)),
            'date': current_date,
            'type2': product.get('type2', '')
        }
        
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # แทนที่ถ้าตรงกับ placeholder
                if cell_text in replacements:
                    cell.text = replacements[cell_text]
                
                # แทนที่ product_type ในช่องที่ merge แล้ว
                elif 'product_type_main' in cell_text:
                    cell.text = cell_text.replace('product_type_main', 
                                                product.get('product_type', ''))
        
    def _merge_files(self, file_paths: List[str], output_path: str):
        """รวมไฟล์หลายไฟล์เป็นไฟล์เดียว"""
        if not file_paths:
            return False
        
        final_doc = Document(file_paths[0])
        
        for i in range(1, len(file_paths)):
            source_doc = Document(file_paths[i])
            for element in source_doc.element.body:
                final_doc.element.body.append(element)
        
        final_doc.save(output_path)
        return True
