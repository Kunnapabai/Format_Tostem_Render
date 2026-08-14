#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# use of main8.py

"""
Site Survey Generator Module
ระบบสร้างไฟล์ Site Survey จากข้อมูล Quotation และ OCR
"""

import os
import re
import tempfile
from datetime import datetime
from typing import Dict, List, Any
from collections import defaultdict
import pandas as pd
from PIL import Image as PILImage
import logging          
import shutil 

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# Document processing
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

# PDF processing
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    print("Warning: ReportLab not available. Install with: pip install reportlab")
    REPORTLAB_AVAILABLE = False

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    print("Warning: pdfplumber not available. Install with: pip install pdfplumber")
    PDF_SUPPORT = False

from quotation_processor import EnhancedQuotationProcessor , EnhancedTOSTEMQuotationProcessor , smart_mosquito_detection_and_merge
from window_door_image_generator import (
    generate_images_for_site_survey,
    get_panel_segments,
    determine_layout,
)


def collapse_doubled_thai(text: str) -> str:
    """ยุบตัวอักษรไทยที่ซ้ำติดกัน - เฉพาะกรณีที่เป็น artifact จากการ extract PDF

    ✅ ยุบ: 'รราาคคาารรววมม' (ทุกตัวซ้ำ = artifact จาก PDF overlay text)
    ❌ ไม่ยุบ: 'รวมมุ้งแล้ว', 'ธรรมดา', 'บรรจุ' (คำไทยที่มีตัวซ้ำจริง)

    เดิมใช้ re.sub(r'([ก-๙])\\1+', r'\\1', text) ตรงๆ ซึ่งทำลายคำไทยที่ถูกต้อง
    จึงตรวจก่อนว่ามีตัวซ้ำหลายจุด (>=3) อันเป็นลักษณะของ artifact จริงๆ
    """
    if not text:
        return text

    doubled_pairs = len(re.findall(r'([ก-๙])\1', text))
    if doubled_pairs < 3:
        return text

    return re.sub(r'([ก-๙])\1+', r'\1', text)


def convert_docx_to_pdf_direct(docx_path: str, pdf_path: str) -> Dict[str, Any]:
    """แปลง DOCX เป็น PDF โดยตรง - รองรับ Cloud/Render Environment"""
    import subprocess
    import os
    
    try:
        # วิธีที่ 1: ใช้ unoconv (แนะนำสำหรับ Linux/Cloud)
        try:
            result = subprocess.run(
                ['unoconv', '--version'],
                capture_output=True, 
                text=True, 
                timeout=5
            )
            
            if result.returncode == 0:
                print("🔧 Using unoconv for PDF conversion...")
                cmd = ['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path]
                
                result = subprocess.run(
                    cmd, 
                    capture_output=True, 
                    text=True, 
                    timeout=120
                )
                
                if result.returncode == 0 and os.path.exists(pdf_path):
                    print(f"✅ แปลง PDF สำเร็จด้วย unoconv: {pdf_path}")
                    return {
                        'success': True,
                        'file_path': pdf_path,
                        'message': 'แปลง DOCX เป็น PDF สำเร็จด้วย unoconv',
                        'method': 'unoconv'
                    }
                else:
                    print(f"⚠️ unoconv failed: {result.stderr}")
                    
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            print(f"unoconv ไม่พร้อมใช้งาน: {e}")
        except Exception as e:
            print(f"unoconv error: {e}")

        # วิธีที่ 2: ใช้ LibreOffice/soffice command line
        try:
            for cmd_name in ['soffice', 'libreoffice']:
                try:
                    result = subprocess.run(
                        [cmd_name, '--version'], 
                        capture_output=True, 
                        text=True, 
                        timeout=5,
                        env={**os.environ, 'HOME': tempfile.gettempdir()}
                    )
                    
                    if result.returncode == 0:
                        print(f"🔧 Using {cmd_name} for PDF conversion...")
                        output_dir = os.path.dirname(pdf_path) or '.'
                        
                        # สร้าง temp user profile
                        temp_profile = tempfile.mkdtemp(prefix='libreoffice_')
                        
                        cmd = [
                            cmd_name,
                            '--headless',
                            '--invisible',
                            '--nodefault',
                            '--nofirststartwizard',
                            '--nolockcheck',
                            '--nologo',
                            '--norestore',
                            f'-env:UserInstallation=file://{temp_profile}',
                            '--convert-to', 'pdf',
                            '--outdir', output_dir,
                            docx_path
                        ]
                        
                        result = subprocess.run(
                            cmd, 
                            capture_output=True, 
                            text=True, 
                            timeout=120,
                            env={**os.environ, 'HOME': temp_profile}
                        )
                        
                        # ลบ temp profile
                        try:
                            shutil.rmtree(temp_profile)
                        except:
                            pass
                        
                        if result.returncode == 0:
                            docx_name = os.path.basename(docx_path)
                            expected_pdf = os.path.join(
                                output_dir, 
                                docx_name.rsplit('.', 1)[0] + '.pdf'
                            )
                            
                            if os.path.exists(expected_pdf):
                                if expected_pdf != pdf_path:
                                    shutil.move(expected_pdf, pdf_path)
                                
                                if os.path.exists(pdf_path):
                                    print(f"✅ แปลง PDF สำเร็จด้วย {cmd_name}: {pdf_path}")
                                    return {
                                        'success': True,
                                        'file_path': pdf_path,
                                        'message': f'แปลง DOCX เป็น PDF สำเร็จด้วย {cmd_name}',
                                        'method': cmd_name
                                    }
                        else:
                            print(f"⚠️ {cmd_name} failed: {result.stderr}")
                            
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    continue
                    
        except Exception as e:
            print(f"LibreOffice/soffice error: {e}")

        # วิธีที่ 3: ใช้ pypandoc
        try:
            import pypandoc
            
            print("🔧 Using pypandoc for PDF conversion...")
            pypandoc.convert_file(
                docx_path,
                'pdf',
                outputfile=pdf_path,
                extra_args=['--pdf-engine=xelatex']
            )
            
            if os.path.exists(pdf_path):
                print(f"✅ แปลง PDF สำเร็จด้วย pypandoc: {pdf_path}")
                return {
                    'success': True,
                    'file_path': pdf_path,
                    'message': 'แปลง DOCX เป็น PDF สำเร็จด้วย pypandoc',
                    'method': 'pypandoc'
                }
                
        except ImportError:
            print("pypandoc ไม่ได้ติดตั้ง")
        except Exception as e:
            print(f"pypandoc error: {e}")

        # วิธีที่ 4: docx2pdf (Windows/Mac)
        try:
            from docx2pdf import convert
            
            print("🔧 Using docx2pdf for PDF conversion...")
            convert(docx_path, pdf_path)
            
            if os.path.exists(pdf_path):
                print(f"✅ แปลง PDF สำเร็จด้วย docx2pdf: {pdf_path}")
                return {
                    'success': True,
                    'file_path': pdf_path,
                    'message': 'แปลง DOCX เป็น PDF สำเร็จด้วย docx2pdf',
                    'method': 'docx2pdf'
                }
                
        except ImportError:
            print("docx2pdf ไม่ได้ติดตั้ง")
        except Exception as e:
            print(f"docx2pdf error: {e}")

        # ❌ ถ้าทุกวิธีล้มเหลว
        error_msg = """
ไม่สามารถแปลง PDF ได้ - กรุณาติดตั้งเครื่องมือใดเครื่องมือหนึ่ง:

สำหรับ Linux/Cloud (Render, Docker):
1. apt-get install libreoffice-writer
2. apt-get install unoconv
3. pip install pypandoc && apt-get install pandoc texlive-xetex

สำหรับ Windows/Mac:
1. ติดตั้ง LibreOffice: https://www.libreoffice.org/
2. pip install docx2pdf (ต้องมี Microsoft Office)
"""
        
        print(f"❌ {error_msg}")
        
        return {
            'success': False,
            'message': error_msg.strip(),
            'error': 'ไม่มีเครื่องมือแปลง PDF ที่ใช้งานได้',
            'suggestions': [
                'apt-get install libreoffice-writer unoconv',
                'pip install pypandoc && apt-get install pandoc texlive-xetex',
                'pip install docx2pdf (Windows/Mac with MS Office)'
            ]
        }
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'message': f'ไม่สามารถแปลง PDF ได้: {str(e)}',
            'error': str(e)
        }
    
def group_products_by_ref(products: List[Dict]) -> Dict[str, List[Dict]]:
    """จัดกลุ่ม products ตาม reference code"""
    grouped = defaultdict(list)
    
    for product in products:
        ref = product.get('ref', '').strip().upper()
        if ref:
            grouped[ref].append(product)
    
    # เรียงลำดับ products ภายในแต่ละกลุ่ม (บานหลักก่อน มุ้งทีหลัง)
    for ref in grouped:
        grouped[ref].sort(key=lambda x: get_product_priority(x))
    
    print(f"Grouped products by ref: {dict(grouped)}")
    return dict(grouped)

def get_product_priority(product: Dict) -> int:
    """กำหนดลำดับความสำคัญของ product (บานหลักก่อน มุ้งทีหลัง)"""
    product_type = product.get('product_type', '').lower()
    
    # ถ้าเป็นมุ้งให้ priority สูง (เรียงหลัง)
    if any(keyword in product_type for keyword in ['มุ้ง', 'mosquito', 'insect', 'screen']):
        return 2
    elif product.get('insect_screen', '').lower() == 'yes':
        return 2
    else:
        return 1
    
def detect_insect_screen_in_group(products_group: List[Dict]) -> str:
    """ตรวจสอบว่าในกลุ่มมีมุ้งหรือไม่"""
    for product in products_group:
        product_type = product.get('product_type', '').lower()
        
        # ตรวจสอบจาก product_type
        mosquito_patterns = [r'\(มุ้ง\)', r'มุ้ง', r'mosquito', r'insect\s*screen', r'screen', r'\(ม\)']
        for pattern in mosquito_patterns:
            if re.search(pattern, product_type, re.IGNORECASE):
                return "Yes"
        
        # ตรวจสอบจาก field เดิม
        if product.get('insect_screen', '').lower() == 'yes':
            return "Yes"
    
    return "No"

def combine_products_in_group(products_group: List[Dict]) -> Dict[str, Any]:
    """รวมข้อมูล products ในกลุ่มเดียวกัน"""
    if not products_group:
        return {}
    
    # ใช้ product แรกเป็นข้อมูลหลัก
    main_product = products_group[0].copy()
    
    # ตรวจสอบมุ้งจากทุก products ในกลุ่ม
    has_mosquito = detect_insect_screen_in_group(products_group)
    main_product['insect_screen'] = has_mosquito
    
    # รวมข้อมูลเพิ่มเติม
    all_remarks = []
    all_types = []
    total_qty = 0
    
    for product in products_group:
        if product.get('product_type'):
            all_types.append(product['product_type'])
        
        if product.get('remarks'):
            all_remarks.append(product['remarks'])
        
        total_qty += product.get('qty', 0)
    
    # อัพเดทข้อมูลที่รวมแล้ว
    if all_types:
        # ใช้ type ที่ไม่ใช่มุ้งเป็นหลัก
        main_types = [t for t in all_types if not any(keyword in t.lower() for keyword in ['มุ้ง', 'mosquito', 'insect'])]
        main_product['product_type'] = main_types[0] if main_types else all_types[0]
    
    if all_remarks:
        main_product['combined_remarks'] = '; '.join(all_remarks)
    
    main_product['total_qty_in_group'] = total_qty
    main_product['products_in_group'] = len(products_group)
    main_product['group_details'] = products_group
    
    return main_product

class EnhancedSiteSurveyGenerator:
    """Site Survey Generator ที่ปรับปรุงแล้ว - แก้ไขการ merge_data และไม่มีหน้าว่าง"""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path or 'site survey.docx'
        # สร้าง quo_processor ใหม่แทนที่จะเป็น None
        self.quo_processor = EnhancedQuotationProcessor()
        self.merged_data = {}
        self.images_by_ref = {}

    def merge_data(self) -> Dict[str, Any]:
        """รวมข้อมูลจาก Quotation และ OCR - แก้ไขปัญหาแล้ว"""
        try:
            print("Starting data merge process...")
            
            # ตรวจสอบและแก้ไขปัญหา quo_processor เป็น None
            if not hasattr(self.quo_processor, 'processed_data'):
                print("Warning: quo_processor.processed_data not found, creating empty data")
                self.quo_processor.processed_data = {
                    'products': [],
                    'project_info': {},
                    'summary': {}
                }
            
            # เริ่มต้นด้วยข้อมูลจาก Quotation
            quo_products = self.quo_processor.processed_data.get('products', [])
            print(f"Quotation products: {len(quo_products)}")
            
            # รวมข้อมูล
            merged_products = []
            for product in quo_products:
                ref = product.get('ref', '')
                
                # รวมข้อมูล
                merged_product = {**product}  # คัดลอกข้อมูลจาก quotation
                
                merged_products.append(merged_product)
            
            # สร้างข้อมูลที่รวมแล้ว
            self.merged_data = {
                'project_info': self.quo_processor.processed_data.get('project_info', {}),
                'products': merged_products,
                'summary': self._calculate_merged_summary(merged_products),
            }
            
            print(f"Merge completed: {len(merged_products)} products processed")
            return self.merged_data
            
        except Exception as e:
            print(f"Error in merge_data: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # Return basic structure even if merge fails
            return {
                'project_info': {},
                'products': [],
                'summary': {},
            }

    def _calculate_merged_summary(self, products: List[Dict]) -> Dict[str, Any]:
        """คำนวดสรุปข้อมูลที่รวมแล้ว"""
        return {
            'total_products': len(products),
            'total_windows': len([p for p in products if p.get('ref', '').startswith('W')]),
            'total_doors': len([p for p in products if p.get('ref', '').startswith('D')]),
            'total_qty': sum(p.get('qty', 0) for p in products),
            'with_insect_screen': len([p for p in products if p.get('insect_screen') == 'Yes'])
        }

    def create_tostem_template(self, output_path: str = None) -> str:
        """สร้าง TOSTEM Site Survey Template ใหม่ - ตาม template เดิม 11 columns"""
        if output_path is None:
            output_path = self.template_path
            
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not available")
                
            doc = Document()
            
            # สร้างตาราง 11 columns (ตาม template เดิม)
            table = doc.add_table(rows=15, cols=11)  # เพิ่มแถวเป็น 15 (เพิ่ม Type3, Type4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # แถวที่ 1: หัวเรื่องหลัก
            self._setup_header_row(table.rows[0])
            
            # แถวที่ 2: ว่าง
            self._merge_row_completely(table.rows[1])
            
            # แถวที่ 3-4: หัวตาราง
            self._setup_table_headers(table.rows[2], table.rows[3])
            
            # แถวที่ 5: ข้อมูลหลัก (product_type_main)
            self._setup_data_row(table.rows[4])
            
            # แถวที่ 6: Type2 + หน่วย mm.
            self._setup_type2_row(table.rows[5])
            
            # แถวที่ 7: Type3
            self._setup_type3_row(table.rows[6])
            
            # แถวที่ 8: Type4
            self._setup_type4_row(table.rows[7])
            
            # แถวที่ 9: ว่าง
            self._merge_row_completely(table.rows[8])
            
            # แถวที่ 10-14: ส่วนล่าง
            self._setup_bottom_sections(table.rows[9:])
            
            doc.save(output_path)
            print(f"สร้าง template สำเร็จที่: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"Error creating template: {e}")
            raise

    def _setup_header_row(self, row):
        """ตั้งค่าแถวหัวเรื่อง"""
        # เซลล์แรกว่าง
        row.cells[0].text = ""
        
        # รวมเซลล์ 1-7 สำหรับ "TOSTEM Site survey WINDOWS & DOORS"
        self._merge_cells_horizontal(row.cells[1], row.cells[7])
        row.cells[1].text = "TOSTEM Site survey WINDOWS & DOORS"
        self._set_cell_bold_center(row.cells[1])
        
        # เซลล์ date
        row.cells[8].text = "date"
        self._set_cell_center(row.cells[8])
        
        # รวมเซลล์ 9-10 สำหรับ "Inside view"
        self._merge_cells_horizontal(row.cells[9], row.cells[10])
        row.cells[9].text = "Inside view"
        self._set_cell_center(row.cells[9])

    def _setup_table_headers(self, row1, row2):
        """ตั้งค่าหัวตาราง - ตาม template เดิม"""
        # Ref (รวม 2 แถว)
        self._merge_cells_vertical(row1.cells[0], row2.cells[0])
        row1.cells[0].text = "Ref"
        self._set_cell_bold_center(row1.cells[0])
        
        # Series (รวม 2 แถว)
        self._merge_cells_vertical(row1.cells[1], row2.cells[1])
        row1.cells[1].text = "Series"
        self._set_cell_bold_center(row1.cells[1])
        
        # Product type (รวม 2 columns, 2 แถว)
        self._merge_cells_horizontal(row1.cells[2], row1.cells[3])
        self._merge_cells_vertical(row1.cells[2], row2.cells[2])
        row1.cells[2].text = "Product type"
        self._set_cell_bold_center(row1.cells[2])
        
        # Opening size
        self._merge_cells_horizontal(row1.cells[4], row1.cells[5])
        row1.cells[4].text = "Opening size"
        self._set_cell_bold_center(row1.cells[4])
        
        # Wo, Ho
        row2.cells[4].text = "Wo"
        row2.cells[5].text = "Ho"
        self._set_cell_center(row2.cells[4])
        self._set_cell_center(row2.cells[5])
        
        # Color (รวม 2 columns 2 แถว)
        self._merge_cells_horizontal(row1.cells[6], row1.cells[7])
        self._merge_cells_vertical(row1.cells[6], row2.cells[6])
        row1.cells[6].text = "Color"
        self._set_cell_bold_center(row1.cells[6])
        
        # Glass (รวม 2 columns 2 แถว)
        self._merge_cells_horizontal(row1.cells[8], row1.cells[9])
        self._merge_cells_vertical(row1.cells[8], row2.cells[8])
        row1.cells[8].text = "Glass"
        self._set_cell_bold_center(row1.cells[8])
        
        # Insect screen (รวม 2 แถว)
        self._merge_cells_vertical(row1.cells[10], row2.cells[10])
        row1.cells[10].text = "Insect\nscreen"
        self._set_cell_bold_center(row1.cells[10])

    def _setup_data_row(self, row):
        """ตั้งค่าแถวข้อมูลหลัก - product_type_main"""
        row.cells[0].text = "ref1"
        row.cells[1].text = "Series1"
        
        # Product type main (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[2], row.cells[3])
        row.cells[2].text = "product_type_main"
        
        row.cells[4].text = "W1"
        row.cells[5].text = "H1"
        
        # Color (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[6], row.cells[7])
        row.cells[6].text = "Color1"
        
        # Glass (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[8], row.cells[9])
        row.cells[8].text = "Glass1"
        
        row.cells[10].text = "Screen1"

    def _setup_type2_row(self, row):
        """ตั้งค่าแถว Type2 (Fixed window F1)"""
        # เซลล์ 0-1 ว่าง
        row.cells[0].text = ""
        row.cells[1].text = ""
        
        # Type2 (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[2], row.cells[3])
        row.cells[2].text = "Type2"
        
        # หน่วย mm.
        row.cells[4].text = "mm."
        row.cells[5].text = "mm."
        
        # เซลล์ 6-10 ว่าง
        for i in range(6, 11):
            row.cells[i].text = ""

    def _setup_type3_row(self, row):
        """ตั้งค่าแถว Type3 (Fixed window F2)"""
        # เซลล์ 0-1 ว่าง
        row.cells[0].text = ""
        row.cells[1].text = ""
        
        # Type3 (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[2], row.cells[3])
        row.cells[2].text = "Type3"
        
        # เซลล์ 4-10 ว่าง
        for i in range(4, 11):
            row.cells[i].text = ""

    def _setup_type4_row(self, row):
        """ตั้งค่าแถว Type4 (Fixed window F3)"""
        # เซลล์ 0-1 ว่าง
        row.cells[0].text = ""
        row.cells[1].text = ""
        
        # Type4 (รวม 2 columns)
        self._merge_cells_horizontal(row.cells[2], row.cells[3])
        row.cells[2].text = "Type4"
        
        # เซลล์ 4-10 ว่าง
        for i in range(4, 11):
            row.cells[i].text = ""



    def _setup_bottom_sections(self, rows):
        """ตั้งค่าส่วนล่างของตาราง"""
        if len(rows) >= 5:
            # แถวที่ 1: More Detail, Note, For giesta series
            header_row = rows[0]
            
            # More Detail (columns 0-2)
            self._merge_cells_horizontal(header_row.cells[0], header_row.cells[2])
            header_row.cells[0].text = "More Detail"
            self._set_cell_bold_center(header_row.cells[0])
            
            # Note (columns 3-6)
            self._merge_cells_horizontal(header_row.cells[3], header_row.cells[6])
            header_row.cells[3].text = "Note"
            self._set_cell_bold_center(header_row.cells[3])
            
            # For giesta series (columns 7-10)
            self._merge_cells_horizontal(header_row.cells[7], header_row.cells[10])
            header_row.cells[7].text = "☐For giesta series"
            self._set_cell_bold_center(header_row.cells[7])
            
            # แถวที่ 2-4: image และ Giesta details (รวม 3 แถว)
            for i in range(1, 4):
                if i < len(rows):
                    # Image (columns 0-2, รวม 3 แถว)
                    if i == 1:
                        self._merge_cells_horizontal(rows[i].cells[0], rows[i].cells[2])
                        if i + 2 < len(rows):
                            self._merge_cells_vertical(rows[i].cells[0], rows[i+2].cells[0])
                        rows[i].cells[0].text = "image"
                        self._set_cell_center(rows[i].cells[0])
                    
                    # Note area (columns 3-6, รวม 3 แถว) - ว่าง
                    if i == 1:
                        self._merge_cells_horizontal(rows[i].cells[3], rows[i].cells[6])
                        if i + 2 < len(rows):
                            self._merge_cells_vertical(rows[i].cells[3], rows[i+2].cells[3])
            
            # แถวที่ 5: ยืนยันการตรวจสอบ
            if len(rows) >= 5:
                confirm_row = rows[4]
                self._merge_cells_horizontal(confirm_row.cells[0], confirm_row.cells[10])
                confirm_row.cells[0].text = "ยืนยันการตรวจสอบ\n(...) (...) (...)\nผู้สรุปผลสำรวจ ผู้ควบคุมงานโครงสร้าง ลูกค้า"
                self._set_cell_center(confirm_row.cells[0])

    # Helper methods สำหรับการจัดการเซลล์
    def _merge_cells_horizontal(self, start_cell, end_cell):
        """รวมเซลล์แนวนอน"""
        try:
            start_cell.merge(end_cell)
        except Exception as e:
            print(f"Error merging horizontal cells: {e}")

    def _merge_cells_vertical(self, top_cell, bottom_cell):
        """รวมเซลล์แนวตั้ง"""
        try:
            top_cell.merge(bottom_cell)
        except Exception as e:
            print(f"Error merging vertical cells: {e}")

    def _merge_row_completely(self, row):
        """รวมทั้งแถว - 11 columns"""
        try:
            self._merge_cells_horizontal(row.cells[0], row.cells[10])
        except Exception as e:
            print(f"Error merging row: {e}")

    def _set_cell_bold_center(self, cell):
        """ตั้งค่าเซลล์ให้เป็นตัวหนาและอยู่กลาง"""
        try:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception as e:
            print(f"Error setting cell format: {e}")

    def _set_cell_center(self, cell):
        """ตั้งค่าเซลล์ให้อยู่กลาง"""
        try:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception as e:
            print(f"Error setting cell center: {e}")

    def _has_meaningful_content(self, doc):
        """ตรวจสอบว่าเอกสารมีเนื้อหาที่มีความหมายหรือไม่ - ปรับปรุงให้เข้มงวดมากขึ้น"""
        try:
            meaningful_data_found = False
            
            # ตรวจสอบข้อความในตาราง
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        
                        # ข้าม header และ template text
                        skip_texts = [
                            '', ' ', 'mm.', 'Ref', 'Series', 'Product type', 
                            'Opening size', 'Wo', 'Ho', 'Color', 'Glass', 
                            'Insect', 'screen', 'More Detail', 'Note', 
                            'For giesta series', 'TOSTEM Site survey WINDOWS & DOORS',
                            'date', 'Inside view', 'Design :', 'Color :', 'Frame type',
                            'Handle', 'Door closer color', 'Door guard color', 
                            'Cylinder lock color', 'For exterior series', 
                            'LATTICE Pattern :', 'Remote control', 'ยืนยันการตรวจสอบ',
                            'ผู้สรุปผลสำรวจ ผู้ควบคุมงานโครงสร้าง ลูกค้า'
                        ]
                        
                        # ข้าม placeholder ที่ไม่ได้ถูกแทนที่
                        placeholder_patterns = [
                            'ref1', 'Series1', 'Color1', 'Glass1', 'Screen1', 
                            'W1', 'H1', 'product_type_main', 'Type2', 'Type3', 'Type4'
                        ]
                        
                        # ข้าม date pattern
                        if re.match(r'\d{4}-\d{2}-\d{2}', cell_text):
                            continue
                            
                        # ข้าม parentheses และ dots
                        if re.match(r'^[\(\)\.]+$', cell_text):
                            continue
                        
                        # ตรวจสอบว่าเป็น meaningful content หรือไม่
                        if (cell_text and 
                            cell_text not in skip_texts and
                            not any(placeholder in cell_text for placeholder in placeholder_patterns) and
                            not cell_text.startswith('Select..') and
                            len(cell_text) > 1 and
                            not cell_text.isdigit()  # ข้าม pure numbers เช่น 2350, 2070
                        ):
                            # ตรวจสอบว่าเป็นข้อมูลจริงที่มีความหมาย
                            if any(keyword in cell_text.lower() for keyword in 
                                   ['door', 'window', 'sliding', 'panel', 'track', 'กระจก', 'มุ้ง']):
                                meaningful_data_found = True
                                print(f"Found meaningful content: '{cell_text}'")
                                break
                
                if meaningful_data_found:
                    break
            
            return meaningful_data_found
            
        except Exception as e:
            print(f"Error checking content: {e}")
            return True  # ถ้าเกิดข้อผิดพลาดให้เก็บไฟล์ไว้

    def _clean_table_content(self, table):
        """ทำความสะอาดเนื้อหาในตาราง - ฟังก์ชันใหม่"""
        try:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    
                    # แทนที่ placeholder ที่เหลือด้วยข้อความว่าง
                    if cell_text in ['Type2','Type3','Type4', 'ref1', 'Series1', 'Color1', 'Glass1', 'Screen1', 'W1', 'H1']:
                        cell.text = ""
                    
                    # ลบ "Select.." text
                    if 'Select..' in cell_text:
                        cell.text = cell_text.replace('Select..', '').strip()
        
        except Exception as e:
            print(f"Error cleaning table: {e}")

    def _remove_empty_content(self, doc):
        """ลบเนื้อหาที่อาจสร้างหน้าว่าง - ปรับปรุงใหม่"""
        try:
            # ตรวจสอบก่อนว่าควรลบหรือไม่
            if not self._has_meaningful_content(doc):
                print("Document appears to be empty or only contains template content")
                return False
            
            # ทำความสะอาดตาราง
            for table in doc.tables:
                self._clean_table_content(table)
                self._remove_empty_rows_from_table(table)
            
            # ลบ paragraph ว่างที่ไม่จำเป็น
            paragraphs_to_remove = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                # ลบเฉพาะ paragraph ที่ว่างเปล่าจริงๆ
                if (not text and 
                    i > 0 and 
                    i < len(doc.paragraphs) - 2):  # เก็บ paragraph แรกและ 2 ตัวสุดท้าย
                    paragraphs_to_remove.append(paragraph)
            
            # ลบ paragraph ที่เลือกไว้
            removed_count = 0
            for paragraph in paragraphs_to_remove[:1]:  # ลบแค่ 1 paragraph เพื่อความปลอดภัย
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                    removed_count += 1
                except:
                    pass
            
            print(f"Cleaned document: removed {removed_count} empty paragraphs")
            return True
                    
        except Exception as e:
            print(f"Warning: Could not clean empty content: {e}")
            return True

    def generate_site_survey_multipage(self, output_path: str) -> Dict[str, Any]:
        """สร้าง Site Survey โดยสร้าง template ใหม่หากไม่มี - แก้ไขปัญหา file path"""
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not available")
            
            # ตรวจสอบว่ามี template หรือไม่
            if not os.path.exists(self.template_path):
                print(f"Template not found: {self.template_path}")
                print("Creating new template...")
                self.create_tostem_template(self.template_path)
            
            # รวมข้อมูล - แก้ไขการเรียก merge_data
            merged_data = self.merge_data()
            products = merged_data.get('products', [])
            
            if not products:
                raise Exception('ไม่มีข้อมูลสินค้าสำหรับสร้างรายงาน')
            
            print(f"Creating Site Survey with {len(products)} products")
            print(f"Output path: {output_path}")
            
            # สร้าง output directory หากไม่มี
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)
            
            # ใช้วิธีแยกไฟล์แล้วรวมกัน - แก้ไขไม่ให้มีหน้าว่าง
            return self._generate_multipage_with_improved_merge(products, merged_data, output_path)
            
        except Exception as e:
            print(f"ERROR: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }
        
    def _generate_multipage_with_improved_merge(self, products, merged_data, output_path):
        """สร้างหลายหน้าแบบรักษา relationships ของรูปภาพ"""
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            import tempfile
            import os
            import shutil
            
            print(f"Processing {len(products)} products...")
            
            if not os.path.exists(self.template_path):
                self.create_tostem_template(self.template_path)
            
            # ✅ วิธีใหม่: สร้างไฟล์แยกแล้ว merge ด้วย python-docx-merge
            temp_files = []
            
            for i, product in enumerate(products):
                try:
                    print(f"\n{'='*60}")
                    print(f"Processing page {i+1} - Ref: {product.get('ref')}")
                    
                    # สร้างไฟล์ชั่วคราว
                    temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
                    os.close(temp_fd)
                    
                    # โหลด template
                    page_doc = Document(self.template_path)
                    
                    # เติมข้อมูล
                    self._fill_template_with_data(page_doc, product, merged_data.get('project_info', {}))
                    
                    # เพิ่มรูปภาพ
                    ref = product.get('ref', '')
                    images = self.images_by_ref.get(ref, [])
                    if images:
                        print(f"Adding {len(images)} images for {ref}")
                        self._add_images_to_template(page_doc, product)
                    else:
                        print(f"No images for {ref}")
                    
                    # บันทึกไฟล์ชั่วคราว
                    page_doc.save(temp_path)
                    temp_files.append(temp_path)
                    
                    print(f"✓ Added page {i+1}")
                    print(f"{'='*60}\n")
                    
                except Exception as e:
                    print(f"✗ Error: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            if not temp_files:
                raise Exception("ไม่มีไฟล์ที่สร้างสำเร็จ")
            
            # ✅ Merge ด้วยวิธีที่รักษา relationships
            success = self._merge_with_relationships(temp_files, output_path)
            
            # ลบไฟล์ชั่วคราว
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except:
                    pass
            
            if success and os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                
                # นับรูปภาพ
                final_doc = Document(output_path)
                actual_images = self._count_images_in_doc(final_doc)
                expected_images = sum(len(imgs) for imgs in self.images_by_ref.values())
                
                print(f"✅ Successfully created: {output_path} ({file_size} bytes)")
                print(f"   Expected images: {expected_images}")
                print(f"   Actual images in file: {actual_images}")
                
                return {
                    'success': True,
                    'file_path': output_path,
                    'message': f'สร้าง Site Survey สำเร็จ ({len(temp_files)} หน้า)',
                    'products_processed': len(temp_files),
                    'total_products': len(products),
                    'file_size': file_size,
                    'images_added': actual_images
                }
            else:
                raise Exception("ไม่สามารถสร้างไฟล์ได้")
                
        except Exception as e:
            print(f"Error: {e}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }

    def _merge_with_relationships(self, file_paths: List[str], output_path: str) -> bool:
        """Merge DOCX files พร้อม relationships (รูปภาพ)"""
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from io import BytesIO
            
            if not file_paths:
                return False
            
            print(f"Merging {len(file_paths)} files with relationships...")
            
            # เริ่มด้วยไฟล์แรก
            final_doc = Document(file_paths[0])
            print(f"  Base: {file_paths[0]}")
            
            # เก็บ image counter
            image_counter = self._count_images_in_doc(final_doc)
            
            # เพิ่มไฟล์อื่นๆ
            for i in range(1, len(file_paths)):
                try:
                    source_doc = Document(file_paths[i])
                    print(f"  Merging file {i+1}: {file_paths[i]}")
                    
                    # สร้าง mapping ระหว่าง old rId กับ new rId
                    rid_mapping = {}
                    
                    # คัดลอก image relationships และสร้าง mapping
                    for rel_id, rel in source_doc.part.rels.items():
                        if "image" in rel.target_ref:
                            try:
                                # ดึงข้อมูลรูปภาพ
                                image_part = rel.target_part
                                image_bytes = image_part.blob
                                image_stream = BytesIO(image_bytes)
                                
                                # เพิ่มรูปภาพและสร้าง relationship ใหม่
                                new_part = final_doc.part.package.image_parts.get_or_add_image_part(image_stream)
                                new_rel_id = final_doc.part.relate_to(new_part, RT.IMAGE)
                                
                                # ✅ relate_to() คืนค่าเป็น string (rId) โดยตรง
                                rid_mapping[rel_id] = new_rel_id
                                
                                image_counter += 1
                                print(f"    ✓ Copied image {image_counter}: {rel_id} -> {new_rel_id}")
                                
                            except Exception as img_error:
                                print(f"    ✗ Image copy error: {img_error}")
                                continue
                    
                    # คัดลอก content และอัพเดท relationship IDs
                    for element in source_doc.element.body:
                        # Clone element
                        new_element = self._deep_copy_element(element)
                        
                        # อัพเดท relationship IDs ในทุก drawing/picture
                        if rid_mapping:
                            self._update_relationship_ids(new_element, rid_mapping)
                        
                        # เพิ่มเข้า final document
                        final_doc.element.body.append(new_element)
                    
                    print(f"  ✓ Merged file {i+1} with {len(rid_mapping)} image mappings")
                    
                except Exception as e:
                    print(f"  ✗ Error: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            # บันทึก
            final_doc.save(output_path)
            
            # ตรวจสอบ
            if os.path.exists(output_path):
                check_doc = Document(output_path)
                final_images = self._count_images_in_doc(check_doc)
                print(f"✅ Merge complete: {final_images} images in final file")
                return True
            
            return False
            
        except Exception as e:
            print(f"Merge error: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _update_relationship_ids(self, element, rid_mapping: dict):
        """อัพเดท relationship IDs ในทุก drawing/picture elements"""
        try:
            from docx.oxml.ns import qn
            
            # ใช้ qn() เพื่อสร้าง qualified name
            blip_tag = qn('a:blip')
            embed_attr = qn('r:embed')
            
            # วนหาทุก blip element ใน tree
            def find_all_blips(elem):
                blips = []
                # ตรวจสอบ element ปัจจุบัน
                if elem.tag == blip_tag:
                    blips.append(elem)
                # ตรวจสอบ children
                for child in elem:
                    blips.extend(find_all_blips(child))
                return blips
            
            # ค้นหาและอัพเดท
            blips = find_all_blips(element)
            updated_count = 0
            
            for blip in blips:
                # ดึง relationship ID เดิม
                old_rid = blip.get(embed_attr)
                
                # ถ้ามี mapping ใหม่ ให้แทนที่
                if old_rid and old_rid in rid_mapping:
                    new_rid = rid_mapping[old_rid]
                    blip.set(embed_attr, new_rid)
                    updated_count += 1
                    print(f"      ✓ Updated rId: {old_rid} -> {new_rid}")
            
            if updated_count == 0:
                print(f"      ℹ️ No blips found to update")
            
        except Exception as e:
            print(f"      ✗ Error updating rIds: {e}")
            import traceback
            traceback.print_exc()

    def _deep_copy_element(self, element):
        """Clone XML element พร้อม attributes"""
        try:
            from copy import deepcopy
            return deepcopy(element)
        except:
            from docx.oxml import parse_xml
            return parse_xml(element.xml)

    def _count_images_in_doc(self, doc) -> int:
        """นับจำนวนรูปภาพใน document"""
        try:
            count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    count += 1
            return count
        except:
            return 0
        
    def _merge_docx_files_without_page_breaks(self, file_paths: List[str], output_path: str) -> bool:
        """รวมไฟล์ DOCX หลายไฟล์โดยไม่มี page break - แก้ไขการ merge รูปภาพ"""
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            if not file_paths:
                return False
            
            print(f"Merging {len(file_paths)} files without page breaks...")
            
            # เริ่มต้นด้วยไฟล์แรก
            final_doc = Document(file_paths[0])
            print(f"  Base file: {file_paths[0]}")
            
            # เพิ่มไฟล์อื่นๆ โดยไม่ใส่ page break
            for i in range(1, len(file_paths)):
                try:
                    source_doc = Document(file_paths[i])
                    print(f"  Adding file {i+1}: {file_paths[i]}")
                    
                    # ✅ วิธีใหม่: คัดลอกเฉพาะ elements ที่ต้องการ
                    for element in source_doc.element.body:
                        # ตรวจสอบว่าเป็น table หรือ paragraph
                        if element.tag.endswith('tbl'):  # Table
                            # Clone table element
                            new_table_element = self._safe_clone_element(element)
                            final_doc.element.body.append(new_table_element)
                            
                        elif element.tag.endswith('p'):  # Paragraph
                            # Clone paragraph element (รวมรูปภาพ)
                            new_para_element = self._safe_clone_element(element)
                            final_doc.element.body.append(new_para_element)
                    
                    print(f"  ✓ Successfully merged file {i+1}")
                    
                except Exception as e:
                    print(f"  ✗ Error merging file {i+1}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            # บันทึกไฟล์สุดท้าย
            final_doc.save(output_path)
            
            # ตรวจสอบผลลัพธ์
            if os.path.exists(output_path) and os.path.getsize(output_path) > 10000:
                print(f"✅ Merge successful: {output_path}")
                
                # ✅ ตรวจสอบว่ามีรูปภาพเท่าไหร่
                merged_doc = Document(output_path)
                image_count = 0
                for rel in merged_doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_count += 1
                print(f"   Total images in merged document: {image_count}")
                
                return True
            else:
                print(f"✗ Merge failed - file size too small or doesn't exist")
                return False
                
        except Exception as e:
            print(f"Error in merge operation: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _safe_clone_element(self, element):
        """Clone XML element อย่างปลอดภัย รวมทั้งรูปภาพ"""
        try:
            from copy import deepcopy
            from docx.oxml import parse_xml
            
            # ใช้ deepcopy แต่ต้องระวังเรื่อง relationship
            cloned = deepcopy(element)
            
            return cloned
            
        except Exception as e:
            print(f"Clone error: {e}, trying alternative method...")
            # Fallback: ใช้ XML string
            from docx.oxml import parse_xml
            return parse_xml(element.xml)
        
    def _merge_docx_files_continuous(self, file_paths: List[str], output_path: str) -> bool:
        """รวมไฟล์ DOCX หลายไฟล์แบบเนื้อหาต่อเนื่อง - ไม่มี page break เลย"""
        try:
            from docx import Document
            
            if not file_paths:
                return False
            
            print(f"Merging {len(file_paths)} files continuously...")
            
            # เริ่มต้นด้วยไฟล์แรก
            final_doc = Document(file_paths[0])
            print(f"  Base file: {file_paths[0]}")
            
            # เพิ่มไฟล์อื่นๆ โดยรวมเนื้อหาต่อเนื่องกัน
            for i in range(1, len(file_paths)):
                try:
                    source_doc = Document(file_paths[i])
                    print(f"  Adding file {i+1}: {file_paths[i]}")
                    
                    # คัดลอกทุก paragraph และ table จาก source
                    for paragraph in source_doc.paragraphs:
                        new_paragraph = final_doc.add_paragraph()
                        new_paragraph.text = paragraph.text
                        # คัดลอก formatting ถ้าต้องการ
                        new_paragraph.style = paragraph.style
                    
                    for table in source_doc.tables:
                        # สร้าง table ใหม่ใน final document
                        new_table = final_doc.add_table(
                            rows=len(table.rows), 
                            cols=len(table.columns)
                        )
                        new_table.style = table.style
                        
                        # คัดลอกข้อมูลใน table
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                new_table.cell(row_idx, col_idx).text = cell.text
                    
                    print(f"  ✓ Successfully merged file {i+1}")
                    
                except Exception as e:
                    print(f"  ✗ Error merging file {i+1}: {e}")
                    continue
            
            # บันทึกไฟล์สุดท้าย
            final_doc.save(output_path)
            
            # ตรวจสอบผลลัพธ์
            if os.path.exists(output_path) and os.path.getsize(output_path) > 10000:
                print(f"✅ Continuous merge successful: {output_path}")
                return True
            else:
                print(f"✗ Merge failed - file size too small or doesn't exist")
                return False
                
        except Exception as e:
            print(f"Error in continuous merge: {e}")
            import traceback
            traceback.print_exc()
            return False
        
    def _deep_copy_element(self, element):
        """Clone XML element พร้อม attributes"""
        try:
            from copy import deepcopy
            return deepcopy(element)
        except:
            from docx.oxml import parse_xml
            return parse_xml(element.xml)
        
    def _clean_glass_text(self, glass_text: str) -> str:
        """ทำความสะอาดข้อความ glass - ลบราคาและข้อความซ้ำออก"""
        if not glass_text:
            return ""
        
        text = glass_text
        
        # ลบข้อความซ้ำซ้อน (รราาคคาา...) - เฉพาะที่เป็น artifact ไม่แตะคำไทยปกติ
        text = collapse_doubled_thai(text)
        
        # ลบข้อความเกี่ยวกับราคา
        text = re.sub(r'ราคา[^\n]*', '', text, flags=re.IGNORECASE)
        
        # ✅ แก้ไขให้ลบเฉพาะราคาที่มีคอมม่า (ตัวเลข 4+ หลัก)
        # ลบรูปแบบ: 22,400.00 หรือ 1,234.56
        text = re.sub(r'\d{1,3}(,\d{3})+\.\d{2}', '', text)
        
        # ✅ เพิ่ม: ลบตัวเลขที่มากกว่า 100 (ไม่มี mm ต่อท้าย)
        # เช่น 8380, 6700 แต่เก็บ 8.38mm, 6mm
        text = re.sub(r'\s+\d{3,}(?!\s*mm)\s*', ' ', text)
        
        # ลบบรรทัดว่าง
        text = re.sub(r'\n+', '\n', text).strip()
        
        # เอาเฉพาะบรรทัดแรก (ข้อมูลกระจก)
        lines = text.split('\n')
        if lines:
            text = lines[0].strip()
        
        return text

    def _finalize_table_layout(self, table):
        """
        ✅ ปรับตารางหลังเติมข้อมูล:
           - เปลี่ยนหัวคอลัมน์ 'Opening size' -> 'Opening Size (mm)'
           - เอาหน่วย 'mm.' ออก (ย้ายไปอยู่ในหัวคอลัมน์แล้ว)
           - ลบแถวว่าง (Type2/Type3/Type4/spacer) ที่ไม่มีข้อมูล เหลือเฉพาะแถวข้อมูล
             (ถ้าแถวไหนมีข้อมูล Fixed sub-panel จะเก็บไว้)
           - ขยายคอลัมน์ Opening Size ไปทางซ้าย โดยลดความกว้าง Product type
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # เซ็ตข้อความโดยคงรูปแบบเดิม (ตัวหนา) ของ run แรก
            def _set_keep_format(cell, text):
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].text = text
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(text)

            # 1) เปลี่ยนหัวคอลัมน์ + หาแถว Wo/Ho และ More Detial
            sub_idx = None   # แถว Wo/Ho
            md_idx = None    # แถว More Detial
            for ri, row in enumerate(table.rows):
                texts = [c.text.strip() for c in row.cells]
                for c in row.cells:
                    if c.text.strip() == 'Opening size':
                        _set_keep_format(c, 'Opening Size (mm)')
                        break
                if any(t == 'Wo' for t in texts):
                    sub_idx = ri
                if md_idx is None and any(t.startswith('More Det') for t in texts):
                    md_idx = ri

            # 2) เคลียร์หน่วย 'mm.'
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() == 'mm.':
                        cell.text = ""

            # 3) ลบแถวว่างระหว่างแถวข้อมูลกับ More Detial (เก็บแถวที่มีข้อมูล Type2/3/4)
            if sub_idx is not None and md_idx is not None:
                data_idx = sub_idx + 1
                rows = table.rows
                to_delete = []
                for ri in range(data_idx + 1, md_idx):
                    if all(not c.text.strip() for c in rows[ri].cells):
                        to_delete.append(rows[ri]._tr)
                for tr in to_delete:
                    tr.getparent().remove(tr)

            # 4) ปรับความกว้างคอลัมน์ (twips) — รวมคงที่ = 10993:
            #    - ขยาย Ref ให้พอสำหรับ Code + floor ยาวสุด เช่น "W13 ชั้น14" (581 -> 1650)
            #    - ขยาย Opening Size ไปทางซ้าย (Product type แคบลง)
            #    - ดึงพื้นที่จาก Product type / Glass / Color (ที่มักสั้น)
            grid_w = [1750, 974, 2183, 378, 967, 936, 530, 470, 1451, 562, 792]
            tbl = table._tbl
            tblPr = tbl.tblPr
            for el in tblPr.findall(qn('w:tblLayout')):
                tblPr.remove(el)
            layout = OxmlElement('w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            tblPr.append(layout)
            grid = tbl.tblGrid
            for c, w in zip(grid.findall(qn('w:gridCol')), grid_w):
                c.set(qn('w:w'), str(w))
            # sync ความกว้างของทุกเซลล์ให้ตรงกับ grid (รองรับ gridSpan)
            for tr in tbl.findall(qn('w:tr')):
                pos = 0
                for tc in tr.findall(qn('w:tc')):
                    tcPr = tc.find(qn('w:tcPr'))
                    span = 1
                    if tcPr is not None:
                        gs = tcPr.find(qn('w:gridSpan'))
                        if gs is not None:
                            span = int(gs.get(qn('w:val')))
                    width = sum(grid_w[pos:pos + span])
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), str(width))
                    tcW.set(qn('w:type'), 'dxa')
                    pos += span

            # 5) จัดแถวลายเซ็นให้ label อยู่กึ่งกลางใต้เส้นประ
            self._fix_signature_alignment(table, sum(grid_w))
        except Exception as e:
            print(f"⚠️ _finalize_table_layout error: {e}")
            import traceback
            traceback.print_exc()

    def _fix_signature_alignment(self, table, table_width_twips):
        """
        ✅ จัดแถวลายเซ็น (ยืนยันการตรวจสอบ) ให้ชื่อผู้เซ็นอยู่กึ่งกลางใต้เส้นประ
           เดิมใช้ช่องว่างจัดเอง ทำให้ label ยาว/สั้นต่างกันเลื่อนไม่ตรงกลาง
           แก้เป็น center tab stop 3 จุด (1/6, 1/2, 5/6) ใช้ทั้งบรรทัดเส้นประและบรรทัดชื่อ
           -> ทุก label อยู่กึ่งกลางตรงกับเส้นประเป๊ะ
        """
        try:
            import re
            from docx.shared import Twips
            from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            stops = [round(table_width_twips * 1 / 6),
                     round(table_width_twips * 1 / 2),
                     round(table_width_twips * 5 / 6)]

            # หา cell ที่มีหัวข้อ 'ยืนยันการตรวจสอบ'
            sig_cell = None
            for row in table.rows:
                for c in row.cells:
                    if 'ยืนยันการตรวจสอบ' in c.text:
                        sig_cell = c
                        break
                if sig_cell:
                    break
            if sig_cell is None:
                return

            for p in sig_cell.paragraphs:
                txt = p.text
                is_line = ('....' in txt) or ('.....)' in txt)
                is_label = ('ผู้สรุป' in txt) or ('ควบคุม' in txt) or ('ลูกค้า' in txt)
                if not (is_line or is_label):
                    continue
                parts = [s.strip() for s in re.split(r'\s{2,}', txt.strip()) if s.strip()]
                if len(parts) != 3:
                    continue

                # เก็บ font เดิม
                fname = p.runs[0].font.name if p.runs else None
                fsize = p.runs[0].font.size if p.runs else None

                # ล้าง runs เดิม
                for r in list(p.runs):
                    r._r.getparent().remove(r._r)

                # ตั้ง center tab stops
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.tab_stops.clear_all()
                for pos in stops:
                    pf.tab_stops.add_tab_stop(Twips(pos), WD_TAB_ALIGNMENT.CENTER)

                # สร้าง run: <tab><item> ต่อกัน 3 ชุด (item จะ center บน tab stop)
                for it in parts:
                    r = p.add_run()
                    r._r.append(OxmlElement('w:tab'))
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = it
                    r._r.append(t)
                    if fname:
                        r.font.name = fname
                    if fsize:
                        r.font.size = fsize
        except Exception as e:
            print(f"⚠️ _fix_signature_alignment error: {e}")
            import traceback
            traceback.print_exc()

    def _fill_template_with_data(self, doc, product: Dict, project_info: Dict):
        """เติมข้อมูลลง template - รองรับ type2, type3, type4 แยกแถว และบวก H สำหรับ Awning + Fixed"""
        try:
            if not doc.tables:
                print("No tables found in template")
                return
            
            # ดึงสีจาก product ก่อน ถ้าไม่มีให้ใช้จาก project_info
            product_color = product.get('color', '') or project_info.get('color', '')
            
            if not product_color:
                print(f"⚠️ Warning: No color found for {product.get('ref', 'Unknown')}")
            else:
                print(f"🎨 Using color: '{product_color}' for {product.get('ref', 'Unknown')}")
            
            # คัดลอก product
            clean_product = product.copy()
            
            # ทำความสะอาดเฉพาะ (Knock-down) และข้อความซ้ำซ้อน
            for field in ["product_type", "Type2", "Type3", "Type4"]:
                if field in clean_product and isinstance(clean_product[field], str):
                    text = clean_product[field]
                    
                    # ลบ (Knock-down)
                    text = re.sub(r'\(\s*Knock-?\s*down\s*\)', '', text, flags=re.IGNORECASE).strip()
                    
                    # ลบตัวอักษรไทยซ้ำ - เฉพาะที่เป็น artifact ไม่แตะคำไทยปกติ
                    text = collapse_doubled_thai(text)
                    
                    # ลบราคาที่อาจติดมา
                    text = re.sub(r'ราคา[^\s]*\s*[\d,\.]+', '', text, flags=re.IGNORECASE).strip()
                    text = re.sub(r'\s+[\d,]+\.?\d*\s*$', '', text).strip()
                    
                    # ลบช่องว่างซ้ำ
                    text = re.sub(r'\s+', ' ', text).strip()
                    
                    if text in ['down)', 'down', ')']:
                        text = ''
                    
                    clean_product[field] = text
            
            # ทำความสะอาด Glass ด้วย
            glass_text = clean_product.get('glass', '')
            if glass_text:
                processor = EnhancedTOSTEMQuotationProcessor()
                clean_glass = processor._clean_glass_text_v2(glass_text)
                clean_product['glass'] = clean_glass
                print(f"  🔍 Cleaned glass: '{glass_text}' → '{clean_glass}'")
            
            table = doc.tables[0]
            current_date = datetime.now().strftime('%Y-%m-%d')
            
            product_type_clean = clean_product.get('product_type', '')
            type2_clean = clean_product.get('Type2', '')
            type3_clean = clean_product.get('Type3', '')
            type4_clean = clean_product.get('Type4', '')
            
            width_value = product.get('width', 0)
            qty_value = product.get('qty', 1)
            calculated_width = width_value * qty_value

            print(f"📐 Calculating width: {width_value} × {qty_value} = {calculated_width}")

            height_value = product.get('height', 0)

            # ✅ บานที่ประกอบจากหลาย product: ใช้ขนาดรวมที่คำนวณไว้แล้ว
            #    (pre_calculate_heights) เพื่อให้ตัวเลขในตารางตรงกับรูปที่วาด
            #    ครอบคลุมทุกชนิด ไม่ใช่แค่ Awning + Fixed
            layout_direction = product.get('layout_direction')

            if layout_direction:
                combo_height = product.get('calculated_height', 0)
                if combo_height > 0:
                    height_value = combo_height

                # ต่อแนวนอนเท่านั้นที่ความกว้างรวมต่างจากความกว้างบานหลัก
                if layout_direction == 'horizontal':
                    combo_width = product.get('calculated_width', 0)
                    if combo_width > 0:
                        calculated_width = combo_width

                print(f"🧩 Combo size from segments ({layout_direction}): "
                      f"{calculated_width} x {height_value}")
            else:
                # 🔥 เดิม: ตรวจสอบและบวก H ถ้าเป็น Awning window + (และมี Type2)
                is_awning_plus = bool(re.search(r'awning\s*window\s*\+', product_type_clean, re.IGNORECASE))
                has_type2 = bool(type2_clean and type2_clean.strip())

                if is_awning_plus and has_type2:
                    type2_height = self._extract_height_from_type2(product, type2_clean)

                    if type2_height > 0:
                        original_height = height_value
                        height_value = original_height + type2_height
                        print(f"🔺 SPECIAL CASE: Awning + Fixed detected")
                        print(f"   Main H: {original_height} + Type2 H: {type2_height} = {height_value}")
                    else:
                        print(f"⚠️ Awning + Fixed detected but no Type2 height found")

            # ✅ เก็บค่าความสูงที่คำนวณแล้วไว้ใน product
            product['calculated_height'] = height_value

            # Mapping ข้อมูล (ใช้ height_value ที่คำนวณแล้ว)
            # ✅ Code = ref + floor (ถ้ามี) เช่น "W1 ชั้น1"
            _code = product.get('ref', '')
            if product.get('floor'):
                _code = f"{_code} {product.get('floor')}"
            replacements = {
                'ref1': _code,
                'Series1': product.get('series', ''),
                'Color1': product_color,
                'Glass1': clean_product.get('glass', ''),
                'Screen1': product.get('insect_screen', 'No'),
                'W1': str(calculated_width),        
                'H1': str(height_value),  # 🔥 ใช้ค่าที่บวกแล้ว (ถ้ามีการบวก)
                'date': current_date,
                'product_type_main': product_type_clean,
                'Type2': type2_clean,
                'Type3': type3_clean,
                'Type4': type4_clean
            }
            
            print(f"=== เติมข้อมูล: {product.get('ref')} ===")
            print(f"  Product Type: '{product_type_clean[:60]}'")
            print(f"  Type2 (F1): '{type2_clean[:80]}'")
            print(f"  Type3 (F2): '{type3_clean[:80]}'")
            print(f"  Type4 (F3): '{type4_clean[:80]}'")
            print(f"  🎨 Color: '{product_color}'")
            print(f"  🔍 Glass: '{clean_product.get('glass', '')[:60]}'")
            print(f"  📏 Final H value: {height_value}")
            
            # แทนที่ในตาราง
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    original_text = cell.text.strip()
                    
                    # ข้ามหัวข้อคอลัมน์
                    if not original_text or original_text in ['Ref', 'Series', 'Product type', 'Color', 'Glass', 'Insect', 'Opening size', 'Wo', 'Ho', 'More Detail', 'Note', '☑For giesta series', 'screen', 'mm.', 'image']:
                        continue
                    
                    # ตรวจสอบและแทนที่ทุก placeholder
                    if original_text in replacements:
                        new_value = replacements[original_text]
                        cell.text = str(new_value) if new_value else ""
                        self._set_cell_font(cell, "TH Sarabun New")

                        # ✅ เติมขนาดของบานย่อยลงในช่อง Wo/Ho ของแถว Type2/3/4
                        if original_text in ('Type2', 'Type3', 'Type4') and new_value:
                            self._fill_sub_panel_size(row, product, original_text)

                        # Debug
                        if new_value:
                            print(f"  ✓ Replaced '{original_text}' with '{str(new_value)[:60]}' at Row {row_index+1}, Cell {cell_index+1}")
            
            # เติมข้อมูลส่วน More Detail และ Note
            self._fill_additional_sections(table, product, project_info)

            # ✅ ปรับ layout ตาราง: หัว 'Opening Size (mm)', ลบแถวว่าง, ปรับความกว้าง
            self._finalize_table_layout(table)

            # เพิ่มรูปภาพ
            self._add_images_to_template(doc, product)

            print(f"=== เสร็จสิ้น: {product.get('ref')} ===\n")
                
        except Exception as e:
            print(f"Error filling template: {str(e)}")
            import traceback
            traceback.print_exc()


    def _fill_sub_panel_size(self, row, product: Dict, type_key: str):
        """
        เติมขนาดของบานย่อย (Fixed window) ลงในช่อง Wo/Ho ของแถว Type2/Type3/Type4

        เดิมช่องนี้เป็นหน่วย 'mm.' ซึ่งย้ายไปอยู่ในหัวคอลัมน์แล้ว จึงว่างอยู่
        """
        try:
            width = product.get(f"{type_key}_W", '')
            height = product.get(f"{type_key}_H", '')

            if not width and not height:
                return

            cells = row.cells

            if len(cells) < 6:
                return

            # ถ้าช่อง Wo/Ho ถูก merge เป็นเซลล์เดียว ให้เขียนรวมกัน
            if cells[4]._tc is cells[5]._tc:
                cells[4].text = f"{width} x {height}"
                self._set_cell_font(cells[4], "TH Sarabun New")
            else:
                cells[4].text = str(width)
                cells[5].text = str(height)
                self._set_cell_font(cells[4], "TH Sarabun New")
                self._set_cell_font(cells[5], "TH Sarabun New")

            print(f"  ✓ {type_key} size: {width} x {height}")

        except Exception as e:
            print(f"  ⚠️ Could not fill {type_key} size: {e}")

    def _extract_height_from_type2(self, product: Dict, type2_text: str) -> int:
        """
        ดึงค่าความสูงของ Fixed window จาก quotation data
        โดยหาจาก merged products (Fixed window ที่ถูก merge เข้ามา)
        """
        try:
            # 🔥 วิธีที่ 1: หาจาก group_details (products ที่ถูก merge เข้ามา)
            if 'group_details' in product:
                for detail in product.get('group_details', []):
                    product_type = detail.get('product_type', '').lower()
                    ref = detail.get('ref', '').upper()
                    
                    # หา Fixed window โดยดูจาก:
                    # 1. ref ที่มี 'F' ต่อท้าย (เช่น D4.0F)
                    # 2. หรือ product_type ที่มี 'fixed'
                    is_fixed_ref = ref.endswith('F') or 'F' in ref[-3:]
                    is_fixed_type = 'fix' in product_type and 'window' in product_type
                    
                    if is_fixed_ref or is_fixed_type:
                        detail_height = detail.get('height', 0)
                        if detail_height > 0:
                            print(f"     ✅ Found Fixed window height from {ref}: {detail_height}mm")
                            return detail_height
            
            # 🔥 วิธีที่ 2: หาจาก fixed_products (ถ้ามีการเก็บแยกไว้)
            if 'fixed_products' in product and isinstance(product['fixed_products'], dict):
                # เช็คทุก Fixed window (F1, F2, F3...)
                for fixed_num, fixed_list in product['fixed_products'].items():
                    if isinstance(fixed_list, list) and len(fixed_list) > 0:
                        first_fixed = fixed_list[0]
                        fixed_height = first_fixed.get('height', 0)
                        if fixed_height > 0:
                            print(f"     ✅ Found Fixed window height from fixed_products: {fixed_height}mm")
                            return fixed_height
            
            # ถ้าไม่พบเลย แสดง warning
            print(f"     ⚠️ Could not find Fixed window height in data")
            return 0
            
        except Exception as e:
            print(f"  ⚠️ Error extracting Type2 height: {e}")
            import traceback
            traceback.print_exc()
            return 0

    def _add_images_to_template(self, doc, product: Dict):
        """เพิ่มรูปภาพลงใน template - ลบข้อความ Image2 แต่ถ้ามีรูปก็ใส่"""
        try:
            if not doc.tables:
                return False
            
            table = doc.tables[0]
            ref = product.get('ref', '')
            images = self.images_by_ref.get(ref, [])
            
            print(f"📸 Adding {len(images)} images for ref: {ref}")
            
            # แยกรูปออกเป็น 2 ประเภท
            panel_images = [img for img in images if img.get('type') == 'panel']
            site_images = [img for img in images if img.get('type') != 'panel']
            
            print(f"   Panel images: {len(panel_images)}, Site images: {len(site_images)}")
            
            # หาตำแหน่งเซลล์
            panel_cell = None  # เซลล์ Image2
            site_image_cell = None  # เซลล์ image
            
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().lower()
                    
                    # 🔥 หาเซลล์ Image2
                    if cell_text == 'image2':
                        panel_cell = cell
                    
                    # หาเซลล์ image
                    elif cell_text == 'image':
                        site_image_cell = cell
            
            images_added = 0
            
            # 1. ใส่รูปบานในเซลล์ Image2 (ถ้ามีรูป)
            if panel_cell:
                # 🔥 ลบข้อความ "Image2" ก่อนเสมอ
                self._clear_cell_content_safely(panel_cell)
                
                if panel_images:
                    # มีรูป → ใส่รูป
                    for img_info in panel_images:
                        img_path = img_info.get('path')
                        
                        if not img_path or not os.path.exists(img_path):
                            continue
                        
                        try:
                            # คำนวณขนาดที่เหมาะสม
                            pil_img = PILImage.open(img_path)
                            width_px, height_px = pil_img.size
                            
                            max_width = 3.5
                            max_height = 3.5
                            aspect_ratio = width_px / height_px
                            
                            if width_px / max_width > height_px / max_height:
                                final_width = max_width
                                final_height = max_width / aspect_ratio
                            else:
                                final_height = max_height
                                final_width = max_height * aspect_ratio
                            
                            para = panel_cell.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            run = para.add_run()
                            run.add_picture(img_path, width=Inches(final_width))
                            
                            images_added += 1
                            print(f"   ✅ Added panel image: {img_info.get('filename')} ({final_width:.2f}x{final_height:.2f} in)")
                        
                        except Exception as e:
                            print(f"   ❌ Error adding panel image: {e}")
                            continue
                else:
                    # ไม่มีรูป → เซลล์ว่าง
                    print(f"   ℹ️  No panel images, Image2 cell is empty")
            
            # 2. ใส่รูป site survey ในเซลล์ image (ถ้ามี)
            # ✅ ลบ placeholder "image" เสมอ แม้ยังไม่มีรูปอัปโหลด (เหมือนเซลล์ Image2)
            #    เดิมลบเฉพาะตอนมีรูป ทำให้เอกสารโชว์คำว่า "image" ค้างไว้
            if site_image_cell:
                self._clear_cell_content_safely(site_image_cell)
                if not site_images:
                    print(f"   ℹ️  No site images, 'image' cell is empty")

            if site_image_cell and site_images:
                for i, img_info in enumerate(site_images):
                    img_path = img_info.get('path')
                    
                    if not img_path or not os.path.exists(img_path):
                        continue
                    
                    try:
                        if i % 2 == 0:
                            para = site_image_cell.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            para = site_image_cell.paragraphs[-1]
                        
                        run = para.add_run()
                        run.add_picture(img_path, width=Inches(2.0) , height=Inches(2.5))
                        
                        if i % 2 == 0 and i < len(site_images) - 1:
                            run.add_text("  ")
                        
                        images_added += 1
                        print(f"   ✅ Added site image: {img_info.get('filename')}")
                        
                    except Exception as e:
                        print(f"   ❌ Error adding site image: {e}")
                        continue
            
            print(f"   ✅ Total images added: {images_added}")
            return images_added > 0
            
        except Exception as e:
            print(f"❌ Error in _add_images_to_template: {str(e)}")
            return False
        
    def _clear_cell_content_safely(self, cell):
        """ลบเนื้อหาใน cell โดยไม่ทำลาย structure"""
        try:
            # เก็บ formatting ไว้
            cell_format = {
                'alignment': cell.vertical_alignment if hasattr(cell, 'vertical_alignment') else None,
                'width': cell.width if hasattr(cell, 'width') else None
            }
            
            # ลบ paragraphs ที่มีอยู่
            for paragraph in cell.paragraphs[1:]:  # เก็บ paragraph แรกไว้
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            
            # ล้าง paragraph แรก
            if cell.paragraphs:
                cell.paragraphs[0].clear()
            
            # คืนค่า formatting
            if cell_format['alignment']:
                cell.vertical_alignment = cell_format['alignment']
                
            return True
            
        except Exception as e:
            print(f"Error clearing cell: {e}")
            return False

    def generate_site_survey_multipage(self, output_path: str) -> Dict[str, Any]:
        """สร้าง Site Survey แบบหลายหน้า - เวอร์ชันที่แก้ไขแล้ว"""
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not available")
            
            # รวมข้อมูล
            merged_data = self.merge_data()
            products = merged_data.get('products', [])
            
            if not products:
                raise Exception('ไม่มีข้อมูลสินค้าสำหรับสร้างรายงาน')
            
            print(f"Creating multipage Site Survey with {len(products)} products")
            
            # สร้าง output directory หากไม่มี
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)
            
            # ใช้ฟังก์ชันที่แก้ไขแล้ว
            return self._generate_multipage_with_improved_merge(products, merged_data, output_path)
            
        except Exception as e:
            print(f"ERROR in generate_site_survey_multipage: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }
    
    def _set_cell_font(self, cell, font_name: str):
        """ตั้งค่าฟอนต์สำหรับเซลล์"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    # สำหรับฟอนต์ไทย ต้องตั้งค่า complex script font ด้วย
                    run.font.cs_name = font_name
                    # ตั้งค่าขนาดฟอนต์ (14pt สำหรับข้อความไทย)
                    run.font.size = Pt(14)
        except Exception as e:
            print(f"Error setting font: {str(e)}")

    def _fill_additional_sections(self, table, product: Dict, project_info: Dict):
        """เติมข้อมูลในส่วน More Detail - Fixed version"""
        try:
            # หาแถวที่มี "More Detail"
            for row_index, row in enumerate(table.rows):
                # Create a list to track which cells we've already processed (merged cells)
                processed_cells = set()
                
                for cell_index, cell in enumerate(row.cells):
                    # Skip if this cell was already processed as part of a merge
                    if id(cell) in processed_cells:
                        continue
                        
                    cell_text = cell.text.strip()
                    
                    # ถ้าเจอเซลล์ More Detail
                    if "More Detail" in cell_text:
                        # เตรียมข้อมูลเพิ่มเติม
                        detail_info = []
                        if product.get('qty', 1) > 1:
                            detail_info.append(f"Quantity: {product.get('qty')} sets")
                        if product.get('remarks'):
                            detail_info.append(f"Remarks: {product.get('remarks')}")
                        if product.get('ocr_matched'):
                            detail_info.append("✓ OCR Data Matched")
                        
                        # Try to find the next cell safely
                        try:
                            # Get all cells in the row
                            all_cells = list(row.cells)
                            
                            # Find current cell position by matching cell text and object
                            current_pos = None
                            for idx, c in enumerate(all_cells):
                                if c.text.strip() == cell_text and id(c) not in processed_cells:
                                    current_pos = idx
                                    break
                            
                            # If found and there's a next cell
                            if current_pos is not None and current_pos + 1 < len(all_cells):
                                next_cell = all_cells[current_pos + 1]
                                next_cell.text = "\n".join(detail_info)
                                # ตั้งค่าฟอนต์สำหรับข้อมูลเพิ่มเติม
                                self._set_cell_font(next_cell, "TH Sarabun New")
                                processed_cells.add(id(next_cell))
                                print(f"✓ Filled More Detail at row {row_index+1}")
                        except Exception as e:
                            print(f"Warning: Could not fill More Detail: {e}")
                    
                    # Mark current cell as processed
                    processed_cells.add(id(cell))
                                    
        except Exception as e:
            print(f"Error filling additional sections: {str(e)}")
            import traceback
            traceback.print_exc()

    def generate_pdf_report(self, output_path: str) -> Dict[str, Any]:
        """สร้างรายงาน PDF โดยแปลงจาก DOCX - ✅ WITH SMART FALLBACK"""
        try:
            # สร้าง DOCX
            docx_path = output_path.replace('.pdf', '_temp.docx')
            docx_result = self.generate_site_survey_multipage(docx_path)
            
            if not docx_result.get('success'):
                return docx_result
            
            # ✅ ตรวจสอบว่า DOCX มีจริง
            if not os.path.exists(docx_path):
                return {
                    'success': False,
                    'message': 'ไม่สามารถสร้างไฟล์ DOCX ได้',
                    'file_path': None
                }
            
            docx_size = os.path.getsize(docx_path)
            logger.info(f"✅ DOCX created: {docx_path} ({docx_size} bytes)")
            
            # ✅ สร้าง backup DOCX ก่อนพยายามแปลง PDF
            docx_backup = output_path.replace('.pdf', '.docx')
            shutil.copy2(docx_path, docx_backup)
            logger.info(f"📋 DOCX backup saved: {docx_backup}")
            
            # พยายามแปลงเป็น PDF
            pdf_result = convert_docx_to_pdf_direct(docx_path, output_path)
            
            # ✅ ตรวจสอบว่า PDF สำเร็จหรือไม่
            if pdf_result.get('success') and os.path.exists(output_path):
                pdf_size = os.path.getsize(output_path)
                
                if pdf_size > 1000:  # PDF ปกติ
                    logger.info(f"✅ PDF created: {output_path} ({pdf_size} bytes)")
                    
                    # ลบ temp DOCX
                    try:
                        os.remove(docx_path)
                    except:
                        pass
                    
                    return {
                        'success': True,
                        'file_path': output_path,
                        'docx_path': docx_backup,  # ✅ เก็บ DOCX ไว้ด้วย
                        'message': 'สร้าง PDF และ DOCX สำเร็จ',
                        'method': pdf_result.get('method'),
                        'file_size': pdf_size,
                        'format': 'pdf'
                    }
            
            # ❌ PDF ล้มเหลว - ใช้ DOCX แทน
            logger.warning(f"⚠️ PDF conversion failed: {pdf_result.get('message')}")
            logger.warning(f"   Falling back to DOCX format")
            
            # ลบ temp DOCX
            try:
                os.remove(docx_path)
            except:
                pass
            
            # ✅ ส่ง DOCX กลับไปแทน (ยังถือว่าสำเร็จ!)
            return {
                'success': True,  # ✅ ยังถือว่าสำเร็จ
                'file_path': docx_backup,
                'message': '✅ สร้าง DOCX สำเร็จ (PDF converter ไม่พร้อมใช้งาน)',
                'format': 'docx',  # ✅ บอกว่าเป็น DOCX
                'file_size': docx_size,
                'pdf_conversion_failed': True,
                'pdf_error': pdf_result.get('message'),
                'suggestions': pdf_result.get('suggestions', [
                    'ติดตั้ง LibreOffice: https://www.libreoffice.org/',
                    'หรือติดตั้ง Microsoft Office (Windows)',
                    'หรือ pip install docx2pdf (Windows + MS Word)'
                ])
            }
            
        except Exception as e:
            logger.error(f"❌ Error in generate_pdf_report: {str(e)}")
            import traceback
            traceback.print_exc()
            
            return {
                'success': False,
                'file_path': None,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }

    def _create_document_pages(self, products: List[Dict], merged_data: Dict, template_path: str) -> List[str]:
        """สร้างหน้าเอกสารสำหรับแต่ละ product - แยกออกมาเป็นฟังก์ชันเดี่ยว"""
        temp_files = []
        valid_files = []
        
        # Ensure template exists
        if not template_path or not os.path.exists(template_path):
            template_path = 'site survey.docx'
            if not os.path.exists(template_path):
                self.create_tostem_template(template_path)
        
        for i, product in enumerate(products):
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_files.append(temp_file.name)
            temp_file.close()
            
            try:
                # Create document
                product_doc = Document(template_path)
                self._fill_grouped_template(product_doc, product, merged_data.get('project_info', {}))
                
                # Save and validate
                if self._remove_empty_content(product_doc):
                    product_doc.save(temp_file.name)
                    
                    saved_doc = Document(temp_file.name)
                    if self._has_meaningful_content(saved_doc):
                        valid_files.append(temp_file.name)
                        ref = product.get('ref', f'Product {i+1}')
                        print(f"Created page for {ref}")

                ref = product.get('ref', '')
                print(f"\n{'='*60}")
                print(f"Processing page {i+1} - Ref: {ref}")
                print(f"Images available: {len(self.images_by_ref.get(ref, []))}")
                
                # เพิ่มรูปภาพ
                images_added = self._add_images_to_template(product_doc, product)
                print(f"Images actually added: {images_added}")
                print(f"{'='*60}\n")
                        
            except Exception as e:
                print(f"Error creating page for product {i+1}: {e}")
                continue
        
        # Clean up unused temp files
        for temp_file in temp_files:
            if temp_file not in valid_files:
                try:
                    os.unlink(temp_file)
                except:
                    pass
        
        return valid_files

    def generate_grouped_site_survey_report(self, output_dir: str, template_path: str = None) -> Dict[str, Any]:
        """สร้าง Site Survey แบบจัดกลุ่ม - เวอร์ชันที่ปรับปรุงแล้ว"""
        try:
            # Get and validate data
            merged_data = self.merge_data()
            products = merged_data.get('products', [])
            
            if not products:
                return {
                    'success': False,
                    'message': 'ไม่มีข้อมูลสินค้าสำหรับสร้างรายงาน'
                }
            
            # Group and combine products
            grouped_products = group_products_by_ref(products)
            combined_products = []
            
            for ref, products_group in grouped_products.items():
                combined_product = combine_products_in_group(products_group)
                if combined_product:
                    combined_products.append(combined_product)
                    mosquito_status = "with mosquito net" if combined_product.get('insect_screen') == 'Yes' else "no mosquito net"
                    print(f"Combined {ref}: {len(products_group)} items -> {mosquito_status}")
            
            # Create document pages
            valid_files = self._create_document_pages(combined_products, merged_data, template_path)
            
            if not valid_files:
                raise Exception("ไม่มีไฟล์ที่มีข้อมูลถูกต้องสำหรับสร้างรายงาน")
            
            # Generate output file
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(output_dir, f'grouped_site_survey_{timestamp}.docx')
            
            # Merge files
            merge_grouped_docx_files(valid_files, output_path)
            
            # Clean up
            for temp_file in valid_files:
                try:
                    os.unlink(temp_file)
                except:
                    pass
            
            return {
                'success': True,
                'file_path': output_path,
                'message': f'สร้าง Grouped Site Survey สำเร็จ ({len(valid_files)} references)',
                'references_processed': len(valid_files),
                'total_products_combined': len(products)
            }
            
        except Exception as e:
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }

    def _validate_and_create_output_file(self, output_path: str, creation_func, *args) -> Dict[str, Any]:
        """Helper function สำหรับสร้างไฟล์และ validate - ลดการเขียนโค้ดซ้ำ"""
        try:
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Execute creation function
            result = creation_func(*args)
            
            # Validate result
            if result.get('success') and os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                if file_size > 1000:  # Minimum reasonable file size
                    result['file_size'] = file_size
                    return result
            
            return {
                'success': False,
                'message': 'ไฟล์ที่สร้างขึ้นไม่ถูกต้องหรือมีขนาดเล็กเกินไป'
            }
            
        except Exception as e:
            return {
                'success': False,
                'message': f'เกิดข้อผิดพลาด: {str(e)}'
            }

    def _fill_grouped_template(self, doc, product: Dict, project_info: Dict):
        """เติมข้อมูลลง template สำหรับ product ที่รวมแล้ว"""
        try:
            if not doc.tables:
                return
            
            table = doc.tables[0]
            current_date = datetime.now().strftime('%Y-%m-%d')
            
            # ✅ Code = ref + floor (ถ้ามี) เช่น "W1 ชั้น1"
            _code = product.get('ref', '')
            if product.get('floor'):
                _code = f"{_code} {product.get('floor')}"
            replacements = {
                'ref1': _code,
                'Series1': product.get('series', ''),
                'Color1': product.get('color', ''),
                'Glass1': product.get('glass', ''),
                'Screen1': product.get('insect_screen', 'No'),
                'W1': str(product.get('width', 0)),
                'H1': str(product.get('height', 0)),
                'date': current_date
            }
            
            product_type_data = product.get('product_type', '')
            
            print(f"=== เติมข้อมูลสำหรับ Grouped Product: {product.get('ref', 'Unknown')} ===")
            print(f"🔍 Insect Screen: {product.get('insect_screen')}")
            print(f"📦 Products in group: {product.get('products_in_group', 1)}")
            
            # แทนที่ข้อมูลในตาราง
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    original_text = cell.text.strip()
                    
                    if not original_text:
                        continue
                    
                    # ข้ามหัวข้อคอลัมน์
                    if original_text in ['Ref', 'Series', 'Product type', 'Color', 'Glass', 'Insect', 'Opening size', 'Wo', 'Ho']:
                        continue
                    
                    new_text = original_text
                    for placeholder, value in replacements.items():
                        if placeholder == original_text:
                            new_text = str(value)
                            print(f"  ✓ Row {row_index+1}, Cell {cell_index+1}: '{placeholder}' → '{value}'")
                            break
                        elif placeholder in original_text:
                            new_text = original_text.replace(placeholder, str(value))
                            break
                    
                    # แทนที่ Product type
                    if original_text == 'product_type_main':
                        new_text = product_type_data
                        print(f"  ✓ Product Type → '{product_type_data}'")
                    
                    if new_text != original_text:
                        cell.text = new_text
            
            # เพิ่มข้อมูลเพิ่มเติม
            try:
                self._fill_grouped_additional_info(table, product, project_info)
            except Exception as e:
                print(f"Warning: Could not fill additional info: {e}")
            
        except Exception as e:
            print(f"เกิดข้อผิดพลาดในการเติมข้อมูล template: {str(e)}")

    def _fill_grouped_additional_info(self, table, product: Dict, project_info: Dict):
        """เติมข้อมูลเพิ่มเติมในส่วน More Detail และ Note"""
        try:
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    try:
                        cell_text = cell.text.strip()
                        
                        if "More Detail" in cell_text:
                            # ใช้การเติมข้อมูลในเซลล์เดียวกัน
                            detail_info = []
                            
                            products_count = product.get('products_in_group', 1)
                            if products_count > 1:
                                detail_info.append(f"Combined: {products_count} items")
                            
                            if product.get('insect_screen') == 'Yes':
                                detail_info.append("✓ Mosquito Net included")
                            
                            if detail_info:
                                cell.text = f"More Detail\n{chr(10).join(detail_info)}"
                                print("✓ Filled More Detail section")
                        
                        elif "Note" in cell_text:
                            notes = []
                            
                            if project_info.get('project_name'):
                                notes.append(f"Project: {project_info.get('project_name')}")
                            
                            if product.get('combined_remarks'):
                                notes.append(f"Combined remarks: {product.get('combined_remarks')}")
                            
                            if notes:
                                cell.text = f"Note\n{chr(10).join(notes)}"
                                print("✓ Filled Note section")
                    
                    except Exception as cell_error:
                        continue
                        
        except Exception as e:
            print(f"Error in additional info: {str(e)}")

def merge_grouped_docx_files(file_paths: List[str], output_path: str):
    """รวมไฟล์ DOCX หลายไฟล์โดยไม่มี page break - เนื้อหาต่อเนื่องกัน"""
    try:
        if not file_paths:
            return False
        
        final_doc = Document(file_paths[0])
        
        for i in range(1, len(file_paths)):
            source_doc = Document(file_paths[i])
            for element in source_doc.element.body:
                final_doc.element.body.append(element)
        
        final_doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error merging files: {e}")
        return False


# ============================================================================
# MAIN FUNCTION
# ============================================================================

def enhanced_process_quotation_file_with_smart_mosquito(file_path: str, start_page: int = 1) -> Dict[str, Any]:
    """
    ประมวลผลไฟล์ Quotation พร้อม Smart Mosquito Detection
    รองรับไฟล์ .xlsx, .xls, .csv, .pdf
    """
    try:
        print(f"📂 Processing quotation file with Smart Mosquito Detection: {file_path}")
        
        # ✅ ตรวจสอบว่าไฟล์มีจริง
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"ไม่พบไฟล์: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f"📄 File type: {file_ext}")
        
        processor = EnhancedQuotationProcessor()
        
        # ประมวลผลไฟล์ตามประเภท
        if file_ext == '.pdf':
            # สำหรับไฟล์ PDF ใช้ TOSTEM processor
            if not PDF_SUPPORT:
                raise ValueError("ไม่รองรับไฟล์ PDF (ขาด pdfplumber library)")
            
            print("🔄 Processing PDF with TOSTEM processor...")
            tostem_processor = EnhancedTOSTEMQuotationProcessor()
            result = tostem_processor.process_tostem_quotation_pdf(file_path)
            
            # Apply smart mosquito detection to PDF results
            if result['success'] and result['data'].get('products'):
                print("🔍 Applying smart mosquito detection to PDF data...")
                original_products = result['data']['products']
                smart_merged_products = smart_mosquito_detection_and_merge(original_products)
                
                # Update result with smart merged products
                result['data']['products'] = smart_merged_products
                result['data']['summary'] = calculate_smart_summary(smart_merged_products)
                
                # Update message
                mosquito_count = len([p for p in smart_merged_products if p.get('insect_screen') == 'Yes'])
                merged_count = len([p for p in smart_merged_products if p.get('merged_from_mosquito')])
                
                result['message'] += f" | Smart Detection: {mosquito_count} with mosquito"
                if merged_count > 0:
                    result['message'] += f" | Auto-merged: {merged_count} refs"
                
                # เก็บข้อมูลใน processor
                processor.processed_data = result['data']
            
            return result
        
        elif file_ext in ['.xlsx', '.xls']:
            # สำหรับไฟล์ Excel
            print("🔄 Processing Excel file...")
            df = pd.read_excel(file_path)
            processor.quo_data = df
            
        elif file_ext == '.csv':
            # สำหรับไฟล์ CSV
            print("🔄 Processing CSV file...")
            df = pd.read_csv(file_path, encoding='utf-8')
            processor.quo_data = df
            
        else:
            raise ValueError(f"ไม่รองรับไฟล์ประเภท {file_ext}")
        
        # ประมวลผลข้อมูลสำหรับไฟล์ Excel/CSV
        print("📊 Extracting products from Excel/CSV...")
        raw_products = processor._extract_enhanced_products(df)
        
        print(f"Found {len(raw_products)} raw products, applying smart mosquito detection...")
        
        # Apply smart mosquito detection
        smart_merged_products = smart_mosquito_detection_and_merge(raw_products)
        
        # สร้างข้อมูลที่ประมวลผลแล้ว
        processor.processed_data = {
            'project_info': processor._extract_project_info(df),
            'products': smart_merged_products,
            'summary': calculate_smart_summary(smart_merged_products)
        }
        
        # สร้างข้อความสรุป
        mosquito_count = len([p for p in smart_merged_products if p.get('insect_screen') == 'Yes'])
        merged_count = len([p for p in smart_merged_products if p.get('merged_from_mosquito')])
        
        message = f'โหลดข้อมูลจาก Quotation สำเร็จ ({len(smart_merged_products)} รายการ)'
        
        if mosquito_count > 0:
            message += f' | มุ้ง: {mosquito_count} รายการ'
        
        if merged_count > 0:
            message += f' | Auto-merged: {merged_count} refs'
        
        print(f"✅ Smart mosquito detection completed: {len(raw_products)} -> {len(smart_merged_products)} products")
        print(f"Products with mosquito: {mosquito_count}, Auto-merged: {merged_count}")
        
        return {
            'success': True,
            'data': processor.processed_data,
            'message': message,
            'smart_detection_stats': {
                'original_count': len(raw_products),
                'final_count': len(smart_merged_products),
                'mosquito_count': mosquito_count,
                'auto_merged_count': merged_count,
                'detection_patterns': [
                    '(มุ้ง)', 'มุ้ง', 'mosquito', 'insect screen', 'net', '(ม)', 'screen'
                ]
            }
        }
        
    except FileNotFoundError as e:
        print(f"❌ File not found error: {str(e)}")
        return {
            'success': False,
            'data': {},
            'message': f'ไม่พบไฟล์: {str(e)}',
            'error_details': str(e)
        }
        
    except ValueError as e:
        print(f"❌ Value error: {str(e)}")
        return {
            'success': False,
            'data': {},
            'message': f'ข้อมูลไม่ถูกต้อง: {str(e)}',
            'error_details': str(e)
        }
        
    except Exception as e:
        print(f"❌ Error in smart mosquito processing: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return {
            'success': False,
            'data': {},
            'message': f'เกิดข้อผิดพลาดในการประมวลผล: {str(e)}',
            'error_details': str(e),
            'traceback': traceback.format_exc()
        }
    
def calculate_smart_summary(products: List[Dict]) -> Dict[str, Any]:
    """คำนวณสรุปข้อมูลสำหรับ smart merged products"""
    try:
        return {
            'total_items': len(products),
            'total_windows': len([p for p in products if p.get('ref', '').startswith('W')]),
            'total_doors': len([p for p in products if p.get('ref', '').startswith('D')]),
            'total_qty': sum(p.get('qty', 0) for p in products),
            'with_mosquito_net': len([p for p in products if p.get('insect_screen') == 'Yes']),
            'without_mosquito_net': len([p for p in products if p.get('insect_screen') == 'No']),
            'auto_merged_refs': len([p for p in products if p.get('merged_from_mosquito', False)]),
            'total_value': sum(p.get('price_total', 0) for p in products),
            'smart_detection_enabled': True
        }
    except Exception as e:
        print(f"Warning: Error calculating summary: {e}")
        return {
            'total_items': len(products),
            'total_windows': 0,
            'total_doors': 0,
            'total_qty': 0,
            'with_mosquito_net': 0,
            'without_mosquito_net': 0,
            'auto_merged_refs': 0,
            'total_value': 0,
            'smart_detection_enabled': True,
            'error': str(e)
        }

def pre_calculate_heights(products: List[Dict]) -> List[Dict]:
    """
    คำนวณขนาดรวมล่วงหน้าสำหรับบานที่ประกอบจากหลาย product
    ก่อนสร้างรูปบาน

    - ต่อแนวตั้ง (กว้างเท่ากัน)  -> H รวม = ผลรวมความสูงของทุกบาน
    - ต่อแนวนอน (สูงเท่ากัน)    -> W รวม = ผลรวมความกว้างของทุกบาน
    - Transom ถูกข้ามไปตั้งแต่ตอน merge จึงไม่ถูกนับ
    """
    try:
        for product in products:
            product_type = product.get('product_type', '').lower()
            height_value = product.get('height', 0)

            # ✅ ทางหลัก: ใช้ panel_segments (รองรับทุกชนิด ไม่ใช่แค่ Awning และรองรับเกิน 2 บาน)
            segments = get_panel_segments(product)

            if len(segments) >= 2:
                direction, total_width, total_height = determine_layout(
                    segments, product.get('ref', '')
                )

                product['calculated_height'] = total_height
                product['calculated_width'] = total_width
                product['original_height'] = height_value
                product['layout_direction'] = direction

                print(f"🔺 Pre-calculated size for {product.get('ref')}: "
                      f"{len(segments)} segments ({direction}) "
                      f"= {total_width}x{total_height} mm")
                continue

            # ⤵️ ทางสำรอง: ข้อมูลเก่าที่ไม่มี panel_segments (Awning + Fixed แบบเดิม)
            # ตรวจสอบ Awning + Fixed
            is_awning_plus = bool(re.search(r'awning\s*window\s*\+', product_type, re.IGNORECASE))

            # ดึง Type2 จาก product
            type2_text = product.get('Type2', '')
            has_type2 = bool(type2_text and type2_text.strip())
            
            if is_awning_plus and has_type2:
                # หาความสูงจาก group_details (Fixed window)
                type2_height = 0
                
                if 'group_details' in product:
                    for detail in product.get('group_details', []):
                        detail_type = detail.get('product_type', '').lower()
                        detail_ref = detail.get('ref', '').upper()
                        
                        is_fixed_ref = detail_ref.endswith('F') or 'F' in detail_ref[-3:]
                        is_fixed_type = 'fix' in detail_type and 'window' in detail_type
                        
                        if is_fixed_ref or is_fixed_type:
                            type2_height = detail.get('height', 0)
                            if type2_height > 0:
                                print(f"  ✅ Found Fixed height for {product.get('ref')}: {type2_height}mm")
                                break
                
                # บวกความสูง
                if type2_height > 0:
                    original_height = height_value
                    calculated_height = original_height + type2_height
                    
                    product['calculated_height'] = calculated_height
                    product['original_height'] = original_height
                    product['type2_height'] = type2_height
                    
                    print(f"🔺 Pre-calculated height for {product.get('ref')}: {original_height} + {type2_height} = {calculated_height}")
                else:
                    print(f"⚠️ Awning + Fixed detected for {product.get('ref')} but no Type2 height")
            else:
                # กรณีปกติ ใช้ความสูงเดิม
                product['calculated_height'] = height_value
        
        return products
        
    except Exception as e:
        print(f"Error in pre_calculate_heights: {e}")
        import traceback
        traceback.print_exc()
        return products
    
def enhanced_generate_site_survey_report(quo_data: Dict,
                                       output_dir: str, template_path: str = None,
                                       images_by_ref: Dict = None,
                                       output_basename: str = None) -> Dict[str, Any]:
    """สร้างรายงาน Site Survey แบบปรับปรุงแล้ว พร้อมรูปบาน

    output_basename: ชื่อไฟล์ (ไม่รวมนามสกุล) ที่ต้องการ เช่น
        'site_survey_คุณภาธร-Quo2026070804'
        ถ้าไม่ระบุ จะใช้ 'enhanced_site_survey_{timestamp}' แบบเดิม
    """
    try:
        print("="*80)
        print("🚀 STARTING ENHANCED SITE SURVEY GENERATION")
        print("="*80)
        
        generator = EnhancedSiteSurveyGenerator(template_path)

        # โหลด site survey images (ถ้ามี)
        if images_by_ref:
            cleaned_images = {}
            for ref, images in images_by_ref.items():
                cleaned_images[ref] = []
                for img in images:
                    img_path = img.get('path') or img.get('filename')
                    if img_path and os.path.exists(img_path):
                        cleaned_images[ref].append({
                            'path': os.path.abspath(img_path),
                            'filename': img.get('filename', os.path.basename(img_path)),
                            'type': 'site'
                        })
            
            generator.images_by_ref = cleaned_images
            print(f"📸 Loaded site images for {len(cleaned_images)} refs")
        else:
            generator.images_by_ref = {}
        
        # โหลดข้อมูล Quotation
        if quo_data and quo_data.get('products'):
            generator.quo_processor.processed_data = quo_data
            print(f"✅ Loaded quotation: {len(quo_data.get('products', []))} products")
        else:
            generator.quo_processor.processed_data = {
                'products': [],
                'project_info': {},
                'summary': {}
            }
        
        merged_data = generator.merge_data()
        products = merged_data.get('products', [])
        
        print(f"📄 Products after merge: {len(products)}")
        
        if not products:
            return {
                'success': False,
                'message': 'ไม่มีข้อมูลสินค้าสำหรับสร้างรายงาน'
            }
        
        # ✅ **คำนวณความสูงก่อนสร้างรูป!**
        print("\n📐 Pre-calculating heights for Awning + Fixed...")
        products = pre_calculate_heights(products)
        
        # อัพเดท products กลับไปใน merged_data
        merged_data['products'] = products

        # ⚠️ generate_site_survey_multipage() เรียก merge_data() ใหม่อีกครั้ง
        #    ซึ่งจะ copy products จาก quo_processor.processed_data
        #    ต้องเขียนค่าที่คำนวณแล้ว (calculated_height / layout_direction) กลับไปด้วย
        #    ไม่งั้นตารางจะใช้ความสูงเดิม แล้วไม่ตรงกับรูปที่วาด
        generator.quo_processor.processed_data['products'] = products
        
        # ตอนนี้ products มี calculated_height แล้ว!
        print(f"✅ Heights calculated for {len(products)} products")
        
        # สร้างรูปบาน (ใช้ calculated_height ✅)
        print("\n🎨 Generating panel images with AI...")
        try:
            panel_images = generate_images_for_site_survey(
                products=products,  
                base_photo_dir="TOSTEM Drawing"
            )
            
            # เพิ่มรูปบานเข้าไปใน images_by_ref
            for ref, panel_image_path in panel_images.items():
                if ref not in generator.images_by_ref:
                    generator.images_by_ref[ref] = []
                
                generator.images_by_ref[ref].insert(0, {
                    'path': str(panel_image_path),
                    'filename': panel_image_path.name,
                    'type': 'panel'
                })
            
            print(f"✅ Generated {len(panel_images)} panel images")
            
        except Exception as e:
            print(f"⚠️ Warning: Could not generate panel images: {e}")
            panel_images = {}
        
        # สร้างไฟล์ - ใช้ชื่อที่ส่งเข้ามา (เช่น site_survey_คุณภาธร-Quo2026070804)
        # ถ้าไม่ได้ส่งมา ใช้ชื่อเดิมที่มี timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = output_basename or f'enhanced_site_survey_{timestamp}'
        print(f"📝 Output basename: {base_name}")

        # สร้าง DOCX
        docx_path = os.path.join(output_dir, f'{base_name}.docx')
        docx_result = generator.generate_site_survey_multipage(docx_path)

        # สร้าง PDF
        pdf_path = os.path.join(output_dir, f'{base_name}.pdf')
        pdf_result = generator.generate_pdf_report(pdf_path)
        
        print("="*80)
        print("✅ COMPLETED")
        print("="*80)
        
        return {
            'success': True,
            'files': {
                'docx': docx_result,
                'pdf': pdf_result
            },
            'merged_data': merged_data,
            'panel_images_generated': len(panel_images),
            'message': f'สร้างรายงานสำเร็จ ({len(products)} หน้า, {len(panel_images)} รูปบาน)'
        }
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'message': f'เกิดข้อผิดพลาด: {str(e)}'
        }
    
def enhanced_generate_grouped_site_survey_report(quo_data: Dict, 
                                               output_dir: str, template_path: str = None) -> Dict[str, Any]:
    """สร้างรายงาน Site Survey แบบจัดกลุ่มรหัสเดียวกัน - ฟังก์ชันหลักใหม่"""
    try:
        print("="*80)
        print("🎯 ENHANCED GROUPED SITE SURVEY GENERATION - NO BLANK PAGES")
        print("="*80)
        
        generator = EnhancedSiteSurveyGenerator(template_path)
        
        # โหลดข้อมูล Quotation
        if quo_data:
            generator.quo_processor.processed_data = quo_data
            print(f"✅ Loaded quotation data: {len(quo_data.get('products', []))} products")
        
        # สร้างรายงานแบบจัดกลุ่ม
        result = generator.generate_grouped_site_survey_report(output_dir, template_path)
        
        print("="*80)
        print("✅ GROUPED SITE SURVEY GENERATION COMPLETED - NO BLANK PAGES")
        print("="*80)
        
        return result
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        return {
            'success': False,
            'message': f'เกิดข้อผิดพลาด: {str(e)}'
        }