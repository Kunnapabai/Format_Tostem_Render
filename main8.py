#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# use of Generate.html

# call function from 3 files - site_survey_generator.py, site_survey_image.py, window_door_image_generator.py, main.py

"""
main.py - Main Flask Application (Refactored & Optimized)
ระบบ OCR + Site Survey Generator ที่แยกโมดูลแล้ว
"""

import os
import json
import uuid
import time
import threading
import shutil
import glob
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any

import re
from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename

# Flask and web
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

from site_survey_generator import (
    enhanced_process_quotation_file_with_smart_mosquito as process_quotation_file,
    enhanced_generate_site_survey_report as generate_site_survey_report,
    DOCX_AVAILABLE,
    REPORTLAB_AVAILABLE
)
from site_survey_image import ImageSupportedSiteSurveyGenerator
from window_door_image_generator import WindowDoorImageGenerator, generate_images_for_site_survey

# เพิ่มตัวแปรนี้
PDF_SUPPORT = True  # หรือ import จาก site_survey ถ้ามี

# Load environment variables
load_dotenv()

# Flask configuration
app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
RESULTS_FOLDER = 'results'
QUOTATION_FOLDER = 'quotations'
SITE_SURVEY_FOLDER = 'site_surveys'
TEMPLATE_FOLDER = 'templates'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}
QUOTATION_EXTENSIONS = {'xlsx', 'xls', 'csv', 'pdf'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create directories
for folder in [UPLOAD_FOLDER, RESULTS_FOLDER, QUOTATION_FOLDER, SITE_SURVEY_FOLDER, TEMPLATE_FOLDER]:
    os.makedirs(folder, exist_ok=True)

SITE_SURVEY_IMAGES_FOLDER = 'site_survey_images'
os.makedirs(SITE_SURVEY_IMAGES_FOLDER, exist_ok=True)

# Job tracking
job_status = {}
job_results = {}
quotation_jobs = {}
site_survey_jobs = {}
site_survey_images = {}
uploaded_images_db = {}

def extract_ref_from_filename(filename):
    """สกัด Ref Code จากชื่อไฟล์"""
    # ลบ extension
    name_without_ext = os.path.splitext(filename)[0]
    
    # Patterns สำหรับหา Ref
    patterns = [
        r'^([DW][A-Z]?\d+(?:\.\d+)?)',           # เริ่มต้น เช่น D1, W2, W1.1
        r'([DW][A-Z]?\d+(?:\.\d+)?)$',           # ลงท้าย เช่น DA1, W2.1
        r'[_\-\s]([DW][A-Z]?\d+(?:\.\d+)?)[_\-\s]',  # อยู่ตรงกลาง
        r'([DW][A-Z]?\d+(?:\.\d+)?)',            # ที่ไหนก็ได้
    ]
    
    for pattern in patterns:
        match = re.search(pattern, name_without_ext, re.IGNORECASE)
        if match:
            return match.group(1).upper()
    
    return None

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def allowed_file(filename, extensions):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in extensions

# ============================================================================
# FLASK ROUTES
# ============================================================================

@app.route('/')
def index():
    """Serve main HTML"""
    return send_from_directory('.', 'index8.html')

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint with Smart Mosquito Detection"""
    try:
        template_exists = os.path.exists('site survey.docx')
            
        return jsonify({
            'status': 'healthy',
            'site_survey_support': True,
            'template_available': template_exists,
            'template_path': 'site survey.docx' if template_exists else None,
            'docx_available': DOCX_AVAILABLE,
            'pdf_support': PDF_SUPPORT,
            'reportlab_available': REPORTLAB_AVAILABLE,
            'smart_mosquito_detection': True,
            'features': {
                'auto_mosquito_merge': True,
                'smart_ref_grouping': True,
                'insect_screen_detection': True
            },
            'message': 'ระบบพร้อมใช้งาน (พร้อม Smart Mosquito Detection)'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'เกิดข้อผิดพลาด: {str(e)}'
        }), 500
    
@app.route('/api/quotation/upload', methods=['POST'])
def upload_quotation():
    """Upload quotation file"""
    if 'file' not in request.files:
        return jsonify({'error': 'ไม่พบไฟล์ในการส่ง'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'ไม่ได้เลือกไฟล์'}), 400
    
    if not allowed_file(file.filename, QUOTATION_EXTENSIONS):
        return jsonify({'error': 'ประเภทไฟล์ไม่ได้รับอนุญาต'}), 400
    
    # รับ start_page จาก form
    start_page = int(request.form.get('start_page', 1))
    
    try:
        job_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        file_path = os.path.join(QUOTATION_FOLDER, f"{job_id}_{filename}")
        file.save(file_path)
        
        # ส่ง start_page ไปยัง processing function
        result = process_quotation_file(file_path, start_page=start_page)
        
        quotation_jobs[job_id] = {
            'status': 'completed' if result['success'] else 'error',
            'original_filename': filename,
            'file_path': file_path,
            'processed_data': result.get('data', {}),
            'message': result['message'],
            'processed_at': datetime.now().isoformat(),
            'start_page': start_page
        }
        
        return jsonify({
            'success': result['success'],
            'job_id': job_id,
            'message': result['message'],
            'start_page': start_page
        })
        
    except Exception as e:
        return jsonify({'error': f'เกิดข้อผิดพลาด: {str(e)}'}), 500
    
# OCR Endpoints
@app.route('/api/upload', methods=['POST'])
def upload_files():
    """Upload and process images for OCR"""
    if 'files' not in request.files:
        return jsonify({'error': 'ไม่พบไฟล์ในการส่ง'}), 400
    
    files = request.files.getlist('files')

    if not files or files[0].filename == '':
        return jsonify({'error': 'ไม่ได้เลือกไฟล์'}), 400
    
    job_id = str(uuid.uuid4())
    
    # Save uploaded files
    uploaded_files = []
    job_folder = os.path.join(UPLOAD_FOLDER, job_id)
    os.makedirs(job_folder, exist_ok=True)
    
    for file in files:
        if file and allowed_file(file.filename, ALLOWED_EXTENSIONS):
            filename = secure_filename(file.filename)
            name, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_filename = f"{name}_{timestamp}{ext}"
            
            file_path = os.path.join(job_folder, unique_filename)
            file.save(file_path)
            uploaded_files.append(file_path)
    
    if not uploaded_files:
        return jsonify({'error': 'ไม่มีไฟล์รูปภาพที่ถูกต้อง'}), 400
    
    # Start async processing
    thread = threading.Thread(
        args=(job_id, uploaded_files, 
              job_status, job_results, RESULTS_FOLDER)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'job_id': job_id,
        'message': 'เริ่มประมวลผลด้วย Gemini AI แล้ว',
        'files_count': len(uploaded_files),
        'ai_service': 'gemini_only'
    })

@app.route('/api/status/<job_id>', methods=['GET'])
def get_job_status(job_id):
    """Get job processing status"""
    if job_id not in job_status:
        return jsonify({'error': 'ไม่พบงานที่ระบุ'}), 404
    return jsonify(job_status[job_id])

@app.route('/api/results/<job_id>', methods=['GET'])
def get_job_results(job_id):
    """Get job processing results"""
    if job_id not in job_results:
        return jsonify({'error': 'ไม่พบผลลัพธ์'}), 404
    return jsonify(job_results[job_id])

@app.route('/api/download/<job_id>', methods=['GET'])
def download_results(job_id):
    """Download OCR results as JSON"""
    filename = f'{job_id}.json'
    file_path = os.path.join(RESULTS_FOLDER, filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'ไม่พบไฟล์ผลลัพธ์'}), 404
    
    return send_from_directory(RESULTS_FOLDER, filename, as_attachment=True)

# Quotation Endpoints
@app.route('/api/quotation/details/<job_id>', methods=['GET'])
def get_quotation_details(job_id):
    """ดูรายละเอียด quotation job รวมทั้งข้อมูลมุ้ง"""
    if job_id not in quotation_jobs:
        return jsonify({'error': 'ไม่พบงานที่ระบุ'}), 404
    
    job = quotation_jobs[job_id]
    
    response = {
        'job_id': job_id,
        'status': job['status'],
        'original_filename': job.get('original_filename'),
        'message': job['message'],
        'processed_at': job['processed_at'],
        'smart_mosquito_enabled': job.get('smart_mosquito_enabled', False),
        'processing_method': job.get('processing_method', 'standard')
    }
    
    # เพิ่มสถิติมุ้ง
    if job.get('mosquito_statistics'):
        response['mosquito_statistics'] = job['mosquito_statistics']
    
    # เพิ่มรายละเอียดสินค้า
    if job['status'] == 'completed' and job.get('processed_data'):
        data = job['processed_data']
        products = data.get('products', [])
        
        response['products_summary'] = []
        mosquito_products = []
        merged_products = []
        
        for product in products:
            product_summary = {
                'ref': product.get('ref'),
                'series': product.get('series'),
                'product_type': product.get('product_type', '')[:100] + ('...' if len(product.get('product_type', '')) > 100 else ''),
                'insect_screen': product.get('insect_screen'),
                'qty': product.get('qty'),
                'merged_from_mosquito': product.get('merged_from_mosquito', False)
            }
            
            # เพิ่มข้อมูลการ merge
            if product.get('merged_from_mosquito'):
                product_summary['merge_info'] = {
                    'mosquito_products_count': product.get('mosquito_products_count', 0),
                    'total_products_merged': product.get('total_products_merged', 0),
                    'combined_remarks': product.get('combined_remarks', '')
                }
                merged_products.append(product_summary)
            
            if product.get('insect_screen') == 'Yes':
                mosquito_products.append(product_summary)
            
            response['products_summary'].append(product_summary)
        
        # เพิ่มสถิติรายละเอียด
        response['detailed_analysis'] = {
            'total_products': len(products),
            'products_with_mosquito': len(mosquito_products),
            'auto_merged_products': len(merged_products),
            'merge_success_rate': f"{(len(merged_products) / max(len(products), 1) * 100):.1f}%",
            'mosquito_detection_patterns_found': [
                '(มุ้ง)', 'มุ้ง', 'mosquito', 'insect screen', 'net', '(ม)', 'screen'
            ]
        }
        
        # เพิ่มสถิติจาก summary
        if data.get('summary'):
            response['statistics'] = data['summary']
    
    return jsonify(response)

# Site Survey Endpoints
@app.route('/api/site-survey/generate', methods=['POST'])
def generate_site_survey():
    """Generate site survey report - Enhanced version"""
    data = request.get_json()

    quotation_job_id = data.get('quotation_job_id')
    
    if not quotation_job_id:
        return jsonify({'error': 'ต้องมีข้อมูล OCR หรือ Quotation อย่างน้อยหนึ่งอย่าง'}), 400
    
    try:
        survey_job_id = str(uuid.uuid4())
        
        # Initialize job status
        site_survey_jobs[survey_job_id] = {
            'status': 'processing',
            'quotation_job_id': quotation_job_id,
            'message': 'กำลังสร้างรายงาน...',
            'generated_at': datetime.now().isoformat()
        }
        
        # Get Quotation data
        quo_data = {}
        if quotation_job_id and quotation_job_id in quotation_jobs:
            quo_data = quotation_jobs[quotation_job_id].get('processed_data', {})
        
        # Find template file
        template_path = 'site survey.docx'
        if not os.path.exists(template_path):
            template_path = None
        
        images_by_ref = {}
        if uploaded_images_db:
            for ref, images in uploaded_images_db.items():
                valid_images = []
                
                for img in images:
                    img_path = img.get('path')
                    
                    if not img_path:
                        app.logger.warning(f"Image has no path: {img.get('filename')}")
                        continue
                    
                    abs_path = os.path.abspath(img_path)
                    
                    if not os.path.exists(abs_path):
                        app.logger.warning(f"Image file not found: {abs_path}")
                        continue

                    valid_images.append({
                        'filename': img['filename'],
                        'path': abs_path,
                        'relative_path': img.get('relative_path', '')
                    })
                    
                    app.logger.info(f"Valid image: {img['filename']} at {abs_path}")
                
                if valid_images:
                    images_by_ref[ref] = valid_images
                    app.logger.info(f"Ref {ref}: {len(valid_images)} valid images")

        # Generate site survey report
        result = generate_site_survey_report(
            quo_data=quo_data,
            output_dir=SITE_SURVEY_FOLDER,
            template_path=template_path,
            images_by_ref=images_by_ref
        )
        
        print(f"Site survey files generated: {result.get('files', {})}")
        for file_type, file_info in result.get('files', {}).items():
            if file_info.get('file_path'):
                print(f"{file_type}: {file_info['file_path']} (exists: {os.path.exists(file_info['file_path'])})")
        
        # Update job with results
        files_info = {}
        if result.get('success') and result.get('files'):
            for file_type, file_result in result['files'].items():
                if file_result.get('success') and file_result.get('file_path'):
                    files_info[file_type] = {
                        'file_path': file_result['file_path'],
                        'success': file_result['success'],
                        'message': file_result.get('message', ''),
                        'size': os.path.getsize(file_result['file_path']) if os.path.exists(file_result['file_path']) else 0
                    }
        
        site_survey_jobs[survey_job_id].update({
            'status': 'completed' if result['success'] else 'error',
            'files': result.get('files', {}),  
            'file_paths': {  
                'docx': result.get('files', {}).get('docx', {}).get('file_path'),
                'pdf': result.get('files', {}).get('pdf', {}).get('file_path')
            },
            'merged_data': result.get('merged_data', {}),
            'message': result['message']
        })
        
        app.logger.info(f"Survey job completed: {survey_job_id}, files: {list(files_info.keys())}")
        
        return jsonify({
            'success': result['success'],
            'survey_job_id': survey_job_id,
            'message': result['message'],
            'files_generated': list(files_info.keys()),
            'products_processed': len(result.get('merged_data', {}).get('products', []))
        })
        
    except Exception as e:
        app.logger.error(f"Error generating site survey: {str(e)}", exc_info=True)
        
        # Update job with error status
        if 'survey_job_id' in locals():
            site_survey_jobs[survey_job_id].update({
                'status': 'error',
                'message': f'เกิดข้อผิดพลาด: {str(e)}',
                'error_details': str(e)
            })
        
        return jsonify({
            'success': False,
            'message': f'เกิดข้อผิดพลาดในการสร้างรายงาน: {str(e)}'
        }), 500

@app.route('/api/site-survey/download/<survey_job_id>/<file_type>', methods=['GET'])
def download_site_survey(survey_job_id, file_type):
    """Download site survey report - Fixed version with better file handling"""
    try:
        print(f"=== DOWNLOAD REQUEST ===")
        print(f"Survey Job ID: {survey_job_id}")
        print(f"File Type: {file_type}")
        print(f"Site Survey Folder: {SITE_SURVEY_FOLDER}")
        
        # ตรวจสอบว่ามี job หรือไม่
        if survey_job_id not in site_survey_jobs:
            print(f"Job not found in site_survey_jobs")
            return jsonify({'error': f'ไม่พบงาน Survey ID: {survey_job_id}'}), 404
        
        job = site_survey_jobs[survey_job_id]
        print(f"Job Status: {job.get('status')}")
        print(f"Job Files: {job.get('files', {})}")
        
        if job.get('status') != 'completed':
            return jsonify({'error': f'งานยังไม่เสร็จสิ้น - สถานะ: {job.get("status")}'}), 400
        
        # วิธีที่ 1: ใช้ข้อมูลจาก job files
        job_files = job.get('files', {})
        if file_type in job_files:
            file_info = job_files[file_type]
            file_path = file_info.get('file_path')
            
            if file_path and os.path.exists(file_path):
                print(f"Found file from job info: {file_path}")
                return send_file(
                    file_path,
                    as_attachment=True,
                    download_name=f'site_survey_{survey_job_id}.{file_type}',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_type == 'docx' else 'application/pdf'
                )
        
        # วิธีที่ 2: ค้นหาไฟล์ในโฟลเดอร์
        print("Searching for files in directory...")
        
        # รายการ pattern ที่จะค้นหา
        search_patterns = [
            f'*{survey_job_id}*.{file_type}',
            f'enhanced_site_survey_*{survey_job_id}*.{file_type}',
            f'site_survey_*{survey_job_id}*.{file_type}',
            f'enhanced_site_survey_*.{file_type}',
            f'site_survey_*.{file_type}'
        ]
        
        import glob
        
        found_files = []
        for pattern in search_patterns:
            search_path = os.path.join(SITE_SURVEY_FOLDER, pattern)
            matching_files = glob.glob(search_path)
            found_files.extend(matching_files)
            print(f"Pattern '{pattern}' found: {len(matching_files)} files")
        
        if found_files:
            # เรียงตามเวลาสร้าง และเอาล่าสุด
            found_files.sort(key=os.path.getctime, reverse=True)
            file_path = found_files[0]
            
            print(f"Found latest file: {file_path}")
            print(f"File size: {os.path.getsize(file_path)} bytes")
            
            return send_file(
                file_path,
                as_attachment=True,
                download_name=f'site_survey_{survey_job_id}.{file_type}',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_type == 'docx' else 'application/pdf'
            )
        
        # วิธีที่ 3: แสดงไฟล์ทั้งหมดในโฟลเดอร์ เพื่อ debug
        print("No files found. Listing all files in directory:")
        try:
            all_files = os.listdir(SITE_SURVEY_FOLDER)
            for f in all_files:
                file_path = os.path.join(SITE_SURVEY_FOLDER, f)
                file_size = os.path.getsize(file_path) if os.path.isfile(file_path) else 0
                print(f"  - {f} ({file_size} bytes)")
        except Exception as e:
            print(f"Error listing directory: {e}")
        
        return jsonify({'error': f'ไม่พบไฟล์ {file_type.upper()}'}), 404
        
    except Exception as e:
        print(f"Download error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'เกิดข้อผิดพลาด: {str(e)}'}), 500

@app.route('/api/jobs/list', methods=['GET'])
def list_jobs():
    """List all jobs"""
    return jsonify({
        'ocr_jobs': {k: {'status': v['status'], 'message': v['message']} for k, v in job_status.items()},
        'quotation_jobs': {k: {'status': v['status'], 'original_filename': v.get('original_filename', '')} for k, v in quotation_jobs.items()},
        'site_survey_jobs': {k: {'status': v['status'], 'message': v['message']} for k, v in site_survey_jobs.items()}
    })

@app.route('/api/cleanup/<job_id>', methods=['DELETE'])
def cleanup_job(job_id):
    """Clean up job data and files"""
    try:
        cleaned = False
        
        # Clean OCR job
        if job_id in job_status:
            job_folder = os.path.join(UPLOAD_FOLDER, job_id)
            if os.path.exists(job_folder):
                shutil.rmtree(job_folder)
            
            result_file = os.path.join(RESULTS_FOLDER, f'{job_id}.json')
            if os.path.exists(result_file):
                os.remove(result_file)
            
            job_status.pop(job_id, None)
            job_results.pop(job_id, None)
            cleaned = True
        
        # Clean quotation job
        if job_id in quotation_jobs:
            job = quotation_jobs[job_id]
            file_path = job.get('file_path')
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
            
            quotation_jobs.pop(job_id, None)
            cleaned = True
        
        # Clean site survey job
        if job_id in site_survey_jobs:
            job = site_survey_jobs[job_id]
            files = job.get('files', {})
            
            for file_info in files.values():
                file_path = file_info.get('file_path')
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
            
            site_survey_jobs.pop(job_id, None)
            cleaned = True
        
        if cleaned:
            return jsonify({'success': True, 'message': 'ลบข้อมูลเรียบร้อย'})
        else:
            return jsonify({'error': 'ไม่พบงานที่ระบุ'}), 404
        
    except Exception as e:
        return jsonify({'error': f'เกิดข้อผิดพลาดในการลบข้อมูล: {str(e)}'}), 500

@app.route('/api/cleanup-all', methods=['DELETE', 'POST'])
def cleanup_all_jobs():
    """ลบข้อมูลทั้งหมดเมื่อรีเฟรชหน้าเว็บ"""
    try:
        cleaned_items = {
            'ocr_jobs': 0,
            'quotation_jobs': 0,
            'site_survey_jobs': 0,
            'upload_folders': 0,
            'result_files': 0,
            'quotation_files': 0,
            'survey_files': 0,
            'uploaded_images': 0
        }
        
        # Clean all OCR jobs
        ocr_job_ids = list(job_status.keys())
        for job_id in ocr_job_ids:
            job_folder = os.path.join(UPLOAD_FOLDER, job_id)
            if os.path.exists(job_folder):
                shutil.rmtree(job_folder)
                cleaned_items['upload_folders'] += 1
            
            result_file = os.path.join(RESULTS_FOLDER, f'{job_id}.json')
            if os.path.exists(result_file):
                os.remove(result_file)
                cleaned_items['result_files'] += 1
            
            cleaned_items['ocr_jobs'] += 1
        
        job_status.clear()
        job_results.clear()
        
        # Clean all quotation jobs
        quotation_job_ids = list(quotation_jobs.keys())
        for job_id in quotation_job_ids:
            job = quotation_jobs[job_id]
            file_path = job.get('file_path')
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
                cleaned_items['quotation_files'] += 1
            
            cleaned_items['quotation_jobs'] += 1
        
        quotation_jobs.clear()
        
        # Clean all site survey jobs
        survey_job_ids = list(site_survey_jobs.keys())
        for job_id in survey_job_ids:
            job = site_survey_jobs[job_id]
            files = job.get('files', {})
            
            for file_info in files.values():
                file_path = file_info.get('file_path')
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
                    cleaned_items['survey_files'] += 1
            
            cleaned_items['site_survey_jobs'] += 1
        
        site_survey_jobs.clear()

        # ลบรูปภาพที่อัปโหลด
        global uploaded_images_db
        image_count = sum(len(images) for images in uploaded_images_db.values())

        for ref, images in uploaded_images_db.items():
            for img in images:
                img_path = img.get('path')
                if img_path and os.path.exists(img_path):
                    try:
                        os.remove(img_path)
                        cleaned_items['uploaded_images'] += 1
                    except:
                        pass

        uploaded_images_db.clear()
        
        # ลบไฟล์เก่าที่อาจตกคั้งในโฟลเดอร์ต่างๆ
        try:
            # ลบไฟล์ทั้งหมดใน UPLOAD_FOLDER (ยกเว้นโฟลเดอร์ที่กำลังใช้งาน)
            for item in os.listdir(UPLOAD_FOLDER):
                item_path = os.path.join(UPLOAD_FOLDER, item)
                try:
                    if os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                    else:
                        os.remove(item_path)
                except:
                    pass
            
            # ลบไฟล์ทั้งหมดใน RESULTS_FOLDER
            for item in os.listdir(RESULTS_FOLDER):
                item_path = os.path.join(RESULTS_FOLDER, item)
                try:
                    os.remove(item_path)
                except:
                    pass
            
            # ลบไฟล์ทั้งหมดใน QUOTATION_FOLDER
            for item in os.listdir(QUOTATION_FOLDER):
                item_path = os.path.join(QUOTATION_FOLDER, item)
                try:
                    os.remove(item_path)
                except:
                    pass
            
            # ลบไฟล์ทั้งหมดใน SITE_SURVEY_FOLDER
            for item in os.listdir(SITE_SURVEY_FOLDER):
                item_path = os.path.join(SITE_SURVEY_FOLDER, item)
                try:
                    os.remove(item_path)
                except:
                    pass

            for item in os.listdir(SITE_SURVEY_IMAGES_FOLDER):
                item_path = os.path.join(SITE_SURVEY_IMAGES_FOLDER, item)
                try:
                    os.remove(item_path)
                except:
                    pass
                    
        except Exception as cleanup_error:
            print(f"Warning during folder cleanup: {cleanup_error}")
        
        return jsonify({
            'success': True,
            'message': 'ลบข้อมูลทั้งหมดสำเร็จ',
            'cleaned_items': cleaned_items
        })
        
    except Exception as e:
        print(f"Error in cleanup_all_jobs: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'เกิดข้อผิดพลาดในการลบข้อมูล: {str(e)}'
        }), 500

# ============================================================================
# IMAGE UPLOAD ENDPOINTS
# ============================================================================

@app.route('/api/site-survey/upload-image', methods=['POST'])
def upload_site_survey_image():
    """อัปโหลดรูปเดียวพร้อมระบุ ref (ถ้ามี)"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'ไม่พบไฟล์'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'ไม่ได้เลือกไฟล์'}), 400
    
    try:
        # รับ ref จาก form (ถ้ามี) หรือสกัดจากชื่อไฟล์
        ref_from_form = request.form.get('ref', '').strip().upper()
        filename = secure_filename(file.filename)
        
        # ลองสกัด ref จากชื่อไฟล์
        ref_from_filename = extract_ref_from_filename(filename)
        
        # ใช้ ref จาก form ก่อน ถ้าไม่มีใช้จากชื่อไฟล์
        final_ref = ref_from_form or ref_from_filename
        
        # สร้างชื่อไฟล์ใหม่ที่ unique
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_ext = os.path.splitext(filename)[1]
        unique_filename = f"{final_ref or 'NOREF'}_{timestamp}_{uuid.uuid4().hex[:8]}{file_ext}"
        
        # บันทึกไฟล์
        file_path = os.path.join(SITE_SURVEY_IMAGES_FOLDER, unique_filename)
        file.save(file_path)

        # ใช้ absolute path ที่แน่นอน
        absolute_path = os.path.abspath(file_path)
        
        # ตรวจสอบว่าไฟล์ถูกสร้างจริง
        if not os.path.exists(absolute_path):
            return jsonify({
                'success': False,
                'error': 'ไม่สามารถบันทึกไฟล์ได้'
            }), 500
        
        # เก็บข้อมูลลง database
        image_data = {
            'filename': filename,
            'stored_filename': unique_filename,
            'path': absolute_path,
            'relative_path': f"{SITE_SURVEY_IMAGES_FOLDER}/{unique_filename}",
            'ref': final_ref,
            'uploaded_at': datetime.now().isoformat(),
            'size': os.path.getsize(absolute_path)
        }
        
        app.logger.info(f"Saved image: {unique_filename}")
        app.logger.info(f"  Absolute path: {absolute_path}")
        app.logger.info(f"  File exists: {os.path.exists(absolute_path)}")
        app.logger.info(f"  File size: {os.path.getsize(absolute_path)} bytes")
    
        # เพิ่มเข้า database
        if final_ref:
            if final_ref not in uploaded_images_db:
                uploaded_images_db[final_ref] = []
            uploaded_images_db[final_ref].append(image_data)
        else:
            if 'UNMATCHED' not in uploaded_images_db:
                uploaded_images_db['UNMATCHED'] = []
            uploaded_images_db['UNMATCHED'].append(image_data)
        
        return jsonify({
            'success': True,
            'message': f'อัปโหลดสำเร็จ: {filename}',
            'ref': final_ref,
            'filename': unique_filename,
            'file_path': image_data['relative_path'],
            'matched': bool(final_ref)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'เกิดข้อผิดพลาด: {str(e)}'
        }), 500

@app.route('/api/site-survey/list-images', methods=['GET'])
def list_all_images():
    """แสดงรูปทั้งหมดแยกตาม ref"""
    try:
        result = {}
        total_count = 0
        
        for ref, images in uploaded_images_db.items():
            result[ref] = {
                'count': len(images),
                'images': [
                    {
                        'filename': img['filename'],
                        'stored_filename': img['stored_filename'],
                        'path': img['relative_path'],
                        'uploaded_at': img['uploaded_at'],
                        'size': img['size']
                    }
                    for img in images
                ]
            }
            total_count += len(images)
        
        return jsonify({
            'success': True,
            'total_images': total_count,
            'refs_count': len(result),
            'images_by_ref': result
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/site-survey/images/<ref>', methods=['GET'])
def get_images_by_ref(ref):
    """ดึงรูปทั้งหมดของ ref ที่ระบุ"""
    try:
        ref = ref.upper()
        
        if ref not in uploaded_images_db:
            return jsonify({
                'success': True,
                'ref': ref,
                'count': 0,
                'images': []
            })
        
        images = uploaded_images_db[ref]
        
        return jsonify({
            'success': True,
            'ref': ref,
            'count': len(images),
            'images': [
                {
                    'filename': img['filename'],
                    'stored_filename': img['stored_filename'],
                    'path': img['relative_path'],
                    'uploaded_at': img['uploaded_at'],
                    'size': img['size']
                }
                for img in images
            ]
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/site-survey/image/<filename>', methods=['DELETE'])
def delete_image(filename):
    """ลบรูปที่ระบุ"""
    try:
        # หารูปใน database
        found = False
        found_ref = None
        found_index = None
        
        for ref, images in uploaded_images_db.items():
            for idx, img in enumerate(images):
                if img['stored_filename'] == filename:
                    found = True
                    found_ref = ref
                    found_index = idx
                    break
            if found:
                break
        
        if not found:
            return jsonify({
                'success': False,
                'error': 'ไม่พบไฟล์ที่ระบุ'
            }), 404
        
        # ลบไฟล์จริง
        image_data = uploaded_images_db[found_ref][found_index]
        if os.path.exists(image_data['path']):
            os.remove(image_data['path'])
        
        # ลบจาก database
        uploaded_images_db[found_ref].pop(found_index)
        
        # ถ้าไม่มีรูปเหลือใน ref นี้ ลบ ref ออก
        if len(uploaded_images_db[found_ref]) == 0:
            del uploaded_images_db[found_ref]
        
        return jsonify({
            'success': True,
            'message': 'ลบรูปสำเร็จ',
            'ref': found_ref,
            'filename': filename
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/site_survey_images/<path:filename>')
def serve_site_survey_image(filename):
    """Serve uploaded site survey images"""
    try:
        return send_from_directory(SITE_SURVEY_IMAGES_FOLDER, filename)
    except Exception as e:
        app.logger.error(f"Image serve error: {e}")
        return jsonify({'error': str(e)}), 404

# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'ไฟล์มีขนาดใหญ่เกินไป (สูงสุด 16MB)'}), 413

@app.errorhandler(500)
def internal_server_error(e):
    return jsonify({'error': 'เกิดข้อผิดพลาดภายในเซิร์ฟเวอร์'}), 500

# ============================================================================
# BACKGROUND TASKS
# ============================================================================

def cleanup_old_files():
    """Clean up old files every 30 minutes"""
    while True:
        try:
            current_time = time.time()
            
            # Clean old upload folders
            for folder_path in glob.glob(os.path.join(UPLOAD_FOLDER, '*')):
                if os.path.isdir(folder_path):
                    folder_age = current_time - os.path.getctime(folder_path)
                    if folder_age > 3600:  # 1 hour
                        shutil.rmtree(folder_path, ignore_errors=True)
            
            # Clean old result files
            for file_path in glob.glob(os.path.join(RESULTS_FOLDER, '*.json')):
                file_age = current_time - os.path.getctime(file_path)
                if file_age > 3600:  # 1 hour
                    os.remove(file_path)
            
            # Clean old quotation files
            for file_path in glob.glob(os.path.join(QUOTATION_FOLDER, '*')):
                file_age = current_time - os.path.getctime(file_path)
                if file_age > 7200:  # 2 hours
                    os.remove(file_path)
            
            # Clean old site survey files
            for file_path in glob.glob(os.path.join(SITE_SURVEY_FOLDER, '*')):
                file_age = current_time - os.path.getctime(file_path)
                if file_age > 7200:  # 2 hours
                    os.remove(file_path)
                    
        except Exception as e:
            print(f"ข้อผิดพลาดในการทำความสะอาด: {e}")
        
        time.sleep(1800)  # 30 minutes

# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_old_files)
cleanup_thread.daemon = True
cleanup_thread.start()

# ============================================================================
# MAIN APPLICATION
# ============================================================================

if __name__ == '__main__':
    print("🚀 Enhanced Site Survey System with Smart Mosquito Detection Starting...")
    print("=" * 80)
    print("📋 Key Features:")
    print("   📸 Multi-page Site Survey (1 ref per page)")
    print("   📸 Enhanced Quotation Processing with Smart Mosquito Detection")
    print("   📸 Automatic Glass & Insect Screen Detection")
    print("   📸 Intelligent Ref Grouping and Auto-Merge")
    print("   📸 OCR Integration with Gemini AI")
    print("   📸 DOCX & PDF Export")
    print("=" * 80)
    print(f"🌐 ใช้งานที่: http://localhost:5000")
    
    # Check dependencies
    print(f"📦 Dependencies Status:")
    print(f"   • python-docx: {'Available' if DOCX_AVAILABLE else 'Not Available'}")
    print(f"   • PDF Support: {'Available' if PDF_SUPPORT else 'Not Available'}")
    print(f"   • ReportLab: {'Available' if REPORTLAB_AVAILABLE else 'Not Available'}")
    
    # Check template file
    template_exists = os.path.exists('site survey.docx')
    print(f"📄 Template Support: {'Enabled' if template_exists else 'Template file not found'}")
    if template_exists:
        print(f"📋 Template Path: site survey.docx")
        try:
            import docx
            doc = docx.Document('site survey.docx')
            print(f"✅ Template loaded successfully ({len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs)")
        except Exception as e:
            print(f"⚠️ Template loading error: {e}")
    
    print("=" * 80)
    print("🔧 Smart Features:")
    print("   • Auto-detect mosquito net from product descriptions")
    print("   • Intelligent merging of same ref with/without mosquito")
    print("   • Enhanced pattern recognition for insect screen")
    print("   • Smart ref grouping and consolidation")
    print("=" * 80)
    
    # Check for critical issues
    missing_deps = []
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx")
    if not REPORTLAB_AVAILABLE:
        missing_deps.append("reportlab")
    
    if missing_deps:
        print(f"⚠️ WARNING: Missing dependencies: {', '.join(missing_deps)}")
        print("   Some features may not work properly.")
    
    if not template_exists:
        print("⚠️ WARNING: Template file 'site survey.docx' not found")
        print("   Site Survey generation will use basic template.")
    
    print("🚀 Starting Flask server with Smart Mosquito Detection...")
    app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)